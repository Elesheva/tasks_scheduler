import os
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import telebot
from telebot import types
import create_db
from apscheduler.schedulers.background import BackgroundScheduler
import sqlite3
from datetime import datetime, timedelta
from docx import Document
bot = telebot.TeleBot('7206218529:AAGXx1IkHVxZ3IrFt09Xgzytanj1n-bpcUI')


# ПРИВЕТСВИЕ И РЕГИСТРАЦИЯ
@bot.message_handler(commands=['start'])
def send_welcome(message):
    create_db.create_db()
    markup = types.InlineKeyboardMarkup()
    markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
    lastname = message.from_user.last_name
    if lastname is None:
        bot.send_message(message.chat.id,
                         f"Здравствуйте, {message.chat.first_name} 👋 Я ваш персональный помощник по планированию задач.\n"
                         f"Чтобы начать работу с ботом необходимо пройти регистрацию."
                         , reply_markup=markup)
    else:
        bot.send_message(message.chat.id,
                         f"Здравствуйте, {message.chat.first_name} {lastname} 👋 Я ваш персональный помощник по планированию задач.\n"
                         f"Чтобы начать работу с ботом необходимо пройти регистрацию."
                         , reply_markup=markup)

@bot.callback_query_handler(func=lambda callback: True)
def registr(callback):
    if callback.data == "registration" or callback.data == "Регистрация":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Пожалуйста, введите свое ФИО")
        bot.register_next_step_handler(callback.message, lambda msg: register_name(msg))
    if callback.data == "Поменять данные" or callback.data == "changing_student":
        bot.send_message(callback.message.chat.id,"Пожалуйста, введите номер поля, которое хотите изменить:")
        bot.register_next_step_handler(callback.message, lambda msg: nomber_change(msg))
    if callback.data == "Поменять данные " or callback.data == "changing_teacher":
        bot.send_message(callback.message.chat.id,"Пожалуйста, введите номер поля, которое хотите изменить:")
        bot.register_next_step_handler(callback.message, lambda msg: nomber_change_teacher(msg))

    if callback.data == "Поменять наименование дисциплины" or callback.data == "changing":
        bot.send_message(callback.message.chat.id,"Пожалуйста, введите номер поля, которое хотите изменить:")
        bot.register_next_step_handler(callback.message, lambda msg: nomber_change_discepline(msg))
    if callback.data == "Добавить дисциплину" or callback.data == "add":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается")
        add_data_to_table_discipline(callback.message)

    if callback.data == "Поменять группу" or callback.data == "changing_group":
        bot.send_message(callback.message.chat.id,"Пожалуйста, введите номер поля, которое хотите изменить:")
        bot.register_next_step_handler(callback.message, lambda msg: nomber_change_group(msg))
    if callback.data == "Добавить группу" or callback.data == "add_group":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Обработка...")
        groap_table(callback.message)

    if callback.data == "delete":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Обработка...")
        statys = 1
        delete_user(callback.message, callback.message.chat.id, statys)
    if callback.data == "deletee":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Обработка...")
        statys = 2
        delete_user(callback.message, callback.message.chat.id, statys)
    if callback.data == "ne_delete":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Принято.")


    if callback.data == "send_completed_task":
        bot.send_message(callback.message.chat.id, "Пожалуйста, введите номер задания, которое хотите отправить:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: send_task_for_teacher(msg, callback.message.chat.id))

    if callback.data == "statystic":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается")
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute(
            "SELECT id, name_of_discipline, group_number, the_task_for_student, send_time, send_date FROM task_for_student WHERE teacher_id = ? AND statys = 1 ",
            (callback.message.chat.id,))
        info_send_task = cursor.fetchall()
        connection.commit()
        connection.close()
        if info_send_task:
            output = "".join(
                f"\n{info_send_task[i][0]}) {info_send_task[i][2]} {info_send_task[i][1]}\nЗадача:\n{info_send_task[i][3]}\nотправлено в {info_send_task[i][4]} {info_send_task[i][5]}"
                for i in
                range(len(info_send_task)))
            bot.send_message(callback.message.chat.id,
                             f"{output}\nПожалуйста, введите номер задания, по которому хотите увидеть статистику:")
            bot.register_next_step_handler(callback.message,
                                           lambda msg: statystics(msg, callback.message.chat.id))
    if callback.data == "send_mark":
        bot.send_message(callback.message.chat.id, "Пожалуйста, введите номер задания, по которому хотите отправить оценку:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: send_comment(msg, callback.message.chat.id))

    if callback.data.startswith("neyd"):
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Записано.")
        perems_str = callback.data[4:]
        perems = perems_str.split('-')
        mark = 2
        bot.send_message(callback.message.chat.id,
                         "Напишите комментарий по выполненной работе:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: ocenka(msg, callback.message.chat.id, mark, perems[0], perems[1]))
    if callback.data.startswith("ydovletvoritelno"):
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Записано.")
        perems_str = callback.data[16:]
        perems = perems_str.split('-')
        mark = 3
        bot.send_message(callback.message.chat.id,
                         "Напишите комментарий по выполненной работе:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: ocenka(msg, callback.message.chat.id, mark, perems[0], perems[1]))
    if callback.data.startswith("horosho"):
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Записано.")
        perems_str = callback.data[7:]
        perems = perems_str.split('-')
        mark = 4
        bot.send_message(callback.message.chat.id,
                         "Напишите комментарий по выполненной работе:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: ocenka(msg, callback.message.chat.id, mark, perems[0], perems[1]))
    if callback.data.startswith("otlichno"):
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Записано.")
        perems_str = callback.data[8:]
        perems = perems_str.split('-')
        mark = 5
        bot.send_message(callback.message.chat.id,
                         "Напишите комментарий по выполненной работе:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: ocenka(msg, callback.message.chat.id, mark, perems[0], perems[1]))

    if callback.data == "change_parol_teacher":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Пожалуйста, введите новый пароль:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: change_parol_teacher(msg, callback.message.chat.id))
    if callback.data == "change_parol_student":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Пожалуйста, введите новый пароль:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: change_parol_student(msg, callback.message.chat.id))

    if callback.data == "done":
        # Обновляем сообщение, убирая кнопки
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш ответ принят.")
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("""                
                        UPDATE statystic_for_student 
                        SET complete = complete + 1, dont_complete = dont_complete - 1
                        WHERE student_id = ?""", (callback.message.chat.id,))
        connection.commit()
        connection.close()

    if callback.data == "dont_done":
        # Обновляем сообщение, убирая кнопки
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш ответ принят.")

    if callback.data == "all_statystic":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute(
            "SELECT name_of_discipline FROM discipline WHERE teacher_id = ?",
            (callback.message.chat.id,))
        discipline = cursor.fetchall()
        connection.commit()
        connection.close()
        if discipline:
            output = "".join(
                f"\n{i+1}) {discipline[i][0]} "
                for i in
                range(len(discipline)))
            bot.send_message(callback.message.chat.id,
                             f"Пожалуйста, выберите дисциплину:\n{output}")
            bot.register_next_step_handler(callback.message,
                                           lambda msg: all_statystic(msg, callback.message.chat.id, discipline))
    if callback.data == "time":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        regular = False
        statys = 1
        send = "time"
        bot.send_message(callback.message.chat.id, "Какую задачу хотите запланировать? Введите название:")
        bot.register_next_step_handler(callback.message, lambda msg: whattime(msg, callback.message.chat.id, regular, statys, send))

    if callback.data == "now":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        regular = False
        statys = 1
        send = "now"
        bot.send_message(callback.message.chat.id, "Какую задачу хотите запланировать? Введите название:")
        bot.register_next_step_handler(callback.message, lambda msg: whattime(msg, callback.message.chat.id, regular, statys, send))
#НАСТРОЙКИ
    if callback.data == "changing_settings_student":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        changing_student(callback.message, callback.message.chat.id)
    if callback.data == "changing_settings_teacher":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        changing_teacher(callback.message, callback.message.chat.id)

    if callback.data == "delete_account":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        delete_zapis(callback.message)

    if callback.data == "all_discipline":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        select_data_for_teacher(callback.message, callback.message.chat.id)

    if callback.data == "all_grupp":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        spisok_grupp(callback.message)

    if callback.data == "info_about_student":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        info(callback.message)

    if callback.data == "change_parol":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        change_parol(callback.message)

    if callback.data == "text_for_admin":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Вы хотите отправить обращение к администратору. В обращении вам нужно подробно изложить сущность вашего обращения. Это поможет администратору быстрее решить проблему. Пожалуйста, введите текст:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: for_admin(msg))

#ЗАДАЧИ
    if callback.data == "add_task":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        text = "add_task"
        new_task(callback.message, text)

    if callback.data == "add_regular_task":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        text = "add_regular_task"
        new_task(callback.message, text)

    if callback.data == "all_tasks":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        get_all_tasks_from_db(callback.message)

    if callback.data == "task_from_the_teacher":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        task_list(callback.message)

    if callback.data == "delete_tasks":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        delete_task_from_db(callback.message)

    if callback.data == "send_markk":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        complete_task(callback.message)

    if callback.data == "send_message_for_student":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        send_message_for_student(callback.message)

    if callback.data == "user_text":
        bot.edit_message_text(chat_id=callback.message.chat.id,
                              message_id=callback.message.message_id,
                              text="Ваш запрос обрабатывается...")
        obrashenya_ot_user(callback.message)

    if callback.data == "otvet_for_user":
        bot.send_message(callback.message.chat.id, "Введите номер обращения:")
        bot.register_next_step_handler(callback.message,
                                       lambda msg: otvet_user(msg))

def register_name(message):
    name = message.text
    bot.send_message(message.chat.id,
                     f"{name}, вы являетесь:\n1. Студентом МУИВ\n2. Преподавателем МУИВ\nВведите номер:")
    bot.register_next_step_handler(message, lambda msg: register_student(msg, name, message.chat.id))

# РЕГИСТРАЦИЯ СТУДЕНТОВ
def register_student(message, name, student_id):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (student_id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (student_id,))
    count = cursor.fetchone()[0]
    if message.text == "1":
        status = 1
        # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
        if count > 0 and teacher == 0:
            bot.send_message(message.chat.id, f"{name}, вы уже зарегистрированы")
            changing_student(message, student_id)
        elif count == 0 and teacher > 0:
            bot.send_message(message.chat.id, f"{name}, вы уже зарегистрированы")
            changing_teacher(message, student_id)
        elif count > 0 and teacher > 0:
            bot.send_message(message.chat.id, f"{name}, вы уже зарегистрированы")
        elif count == 0 and teacher == 0:
            bot.send_message(student_id, "Введите пароль: ")
            bot.register_next_step_handler(message, lambda msg: proverka_parol(msg, name, student_id, status))
    elif message.text == "2":
        status = 2
        # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
        if count > 0 and teacher == 0:
            bot.send_message(message.chat.id, f"{name}, вы уже зарегистрированы")
            changing_student(message, student_id)
        elif count == 0 and teacher > 0:
            bot.send_message(message.chat.id, f"{name}, вы уже зарегистрированы")
            changing_teacher(message, student_id)
        elif count > 0 and teacher > 0:
            bot.send_message(message.chat.id, f"{name}, вы уже зарегистрированы")
        elif count == 0 and teacher == 0:
            bot.send_message(student_id, "Введите пароль: ")
            bot.register_next_step_handler(message, lambda msg: proverka_parol(msg, name, student_id, status))
    else:
        bot.send_message(message.chat.id,
                         f"{name}, вы ввели неверное значение, попробуйте ещё раз.\nВы являетесь:\n1. Студентом МУИВ\n2. Преподавателем МУИВ\nВведите номер:")
        bot.register_next_step_handler(message, lambda msg: register_student(msg, name, student_id))

def proverka_parol(message, name, student_id, status):
    parol = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    if status == 1:
        cursor.execute("SELECT COUNT (*) FROM parol WHERE parol_for_student = ?",
                       (parol,))
        student_parol = cursor.fetchone()[0]
        if student_parol > 0:
            bot.send_message(message.chat.id,
                             f"{name}, вы являетесь студентом МУИВ, пожалуйста введите ваш номер телефона.\nЭти данные будут доступны только вашему преподавателю")
            bot.register_next_step_handler(message, lambda msg: student_nomber(msg, name, student_id))
        else:
            bot.send_message(message.chat.id, "Вы ввели неправильный пароль. Попробуйте ещё раз:")
            bot.register_next_step_handler(message, lambda msg: proverka_parol(msg, name, student_id, status))

    if status == 2:
        cursor.execute("SELECT COUNT (*) FROM parol WHERE parol_for_teacher = ?",
                       (parol,))
        teacher_parol = cursor.fetchone()[0]
        if teacher_parol > 0:
            bot.send_message(message.chat.id,
                             f"{name}, вы являетесь преподавателем МУИВ, пожалуйста введите ваш номер телефона:")
            bot.register_next_step_handler(message, lambda msg: register_teacher(msg, name, message.chat.id))
        else:
            bot.send_message(message.chat.id, "Вы ввели неправильный пароль. Попробуйте ещё раз:")
            bot.register_next_step_handler(message, lambda msg: proverka_parol(msg, name, student_id, status))

def student_nomber(message, name, student_id):
    phone_nomber = message.text
    bot.send_message(message.chat.id, f"{name}, введите вашу почту\nПример: plan_it@mail.com")
    bot.register_next_step_handler(message, lambda msg: mail_student(msg, name, student_id, phone_nomber))


def mail_student(message, name, student_id, phone_nomber):
    mail = message.text
    bot.send_message(message.chat.id, f"{name}, ваш пол:\n1. Мужской\n2. Женский\nВведите цифру с нужным вариантом:")
    bot.register_next_step_handler(message, lambda msg: gender_student(msg, name, student_id, phone_nomber, mail))


def gender_student(message, name, student_id, phone_nomber, mail):
    gender = message.text
    if gender == "1" or gender == "2":
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("SELECT DISTINCT faculty FROM discipline")
        info_about_faculty = cursor.fetchall()
        connection.commit()
        connection.close()
        output = "".join(
            f"{i+1}) {info_about_faculty[i][0]}\n"
            for i in range(len(info_about_faculty)))
        bot.send_message(message.chat.id, f"{name}, укажите ваш Факультет:\n{output} ")
        bot.register_next_step_handler(message,
                                       lambda msg: faculty_student(msg, name, student_id, phone_nomber, mail, gender,
                                                                   info_about_faculty))
    else:
        bot.send_message(message.chat.id,
                         f"{name}, вы ввели неверное значение\nваш пол:\n1. Мужской\n2. Женский\nВведите цифру с нужным вариантом:")
        bot.register_next_step_handler(message, lambda msg: gender_student(msg, name, student_id, phone_nomber, mail))


def faculty_student(message, name, student_id, phone_nomber, mail, gender, info_about_faculty):
    try:
        nomber = int(message.text)
    except ValueError:
        bot.send_message(message.chat.id, "Неверный номер факультета. Пожалуйста, попробуйте снова.")
        bot.register_next_step_handler(message,
                                       lambda msg: faculty_student(msg, name, student_id, phone_nomber, mail, gender,
                                                                   info_about_faculty))
        return
    have = False
    for i in range(len(info_about_faculty)):
            have = True
            faculty = info_about_faculty[i][nomber - 1]
            bot.send_message(message.chat.id, f"{name}, укажите ваш курс\nПример: 1")
            bot.register_next_step_handler(message,
                                           lambda msg: course_student(msg, name, student_id, phone_nomber, mail, gender,
                                                                      faculty))
    if not have:
        bot.send_message(message.chat.id, "Неверный номер факультета. Пожалуйста, попробуйте снова.")
        bot.register_next_step_handler(message,
                                       lambda msg: faculty_student(msg, name, student_id, phone_nomber, mail, gender,
                                                                   info_about_faculty))


def course_student(message, name, student_id, phone_nomber, mail, gender, faculty):
    course = message.text
    if course == "1" or course == "2" or course == "3" or course == "4":
        course = int(message.text)
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("SELECT id, group_number, faculty, course FROM groups WHERE faculty = ? AND course = ? ",
                       (faculty, course))
        info_about_faculty = cursor.fetchall()
        connection.commit()
        connection.close()
        output = "".join(
            f"{info_about_faculty[i][0]}) {info_about_faculty[i][1]}, факультет: {info_about_faculty[i][2]}, курс: {info_about_faculty[i][3]}\n"
            for i in range(len(info_about_faculty)))
        if info_about_faculty:
            bot.send_message(message.chat.id, f"{name}, выберите номер группы:\n{output} ")
            bot.register_next_step_handler(message,
                                           lambda msg: group_number(msg, name, student_id, phone_nomber, mail, gender,
                                                                    faculty, course, info_about_faculty))
        else:
            bot.send_message(message.chat.id, f"{name}, к сожалению, группа пока не создана.")
    else:
        bot.send_message(message.chat.id,
                         f"{name}, вы ввели неверное значение, укажите ваш курс.\nПример: 1")
        bot.register_next_step_handler(message,
                                       lambda msg: course_student(msg, name, student_id, phone_nomber, mail, gender,
                                                                  faculty))


def group_number(message, name, student_id, phone_nomber, mail, gender, faculty, course, info_about_faculty):
    try:
        nomber = int(message.text)
    except ValueError:
        bot.send_message(message.chat.id,
                         f"{name}, вы ввели неверное значение, укажите вашу группу:")
        bot.register_next_step_handler(message,
                                       lambda msg: group_number(msg, name, student_id, phone_nomber, mail, gender,
                                                                faculty, course, info_about_faculty))
        return

    have = False
    for i in range(len(info_about_faculty)):
        if info_about_faculty[i][0] == nomber:
            have = True
            group = info_about_faculty[i][1]
            connection = sqlite3.connect('my_database.db')
            cursor = connection.cursor()
            cursor.execute(
                'INSERT INTO student (student_id, name, phone_number, mail, gender, faculty, course, group_number) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                (student_id, name, phone_nomber, mail, gender, faculty, course, group))
            connection.commit()
            connection.close()
            bot.send_message(message.chat.id, "Вы зарегистрированы!")
            changing_student(message, student_id)
    if not have:
        bot.send_message(message.chat.id,
                         f"{name}, вы ввели неверное значение, укажите вашу группу:")
        bot.register_next_step_handler(message,
                                       lambda msg: group_number(msg, name, student_id, phone_nomber, mail, gender,
                                                                faculty, course, info_about_faculty))


# РЕГИСТРАЦИЯ ПРЕПОДАВАТЕЛЯ
def register_teacher(message, name, teacher_id):
    teacher_phone_nomber = message.text
    bot.send_message(message.chat.id, f"{name}, введите вашу почту\nПример: plan_it@mail.com")
    bot.register_next_step_handler(message, lambda msg: mail_teacher(msg, name, teacher_id, teacher_phone_nomber))


def mail_teacher(message, name, teacher_id, teacher_phone_nomber):
    mail = message.text
    bot.send_message(message.chat.id, f"{name}, ваш пол:\n1. Мужской\n2. Женский\nВведите цифру с нужным вариантом:")
    bot.register_next_step_handler(message,
                                   lambda msg: gender_teacher(msg, name, teacher_id, teacher_phone_nomber, mail))


def gender_teacher(message, name, teacher_id, teacher_phone_nomber, mail):
    gender = message.text
    if gender == "1" or gender == "2":
        bot.send_message(message.chat.id, f"{name}, укажите наименование вашей кафедры\nПример: Информационные системы")
        bot.register_next_step_handler(message,
                                       lambda msg: department_teacher(msg, name, teacher_id, teacher_phone_nomber, mail,
                                                                      gender))
    else:
        bot.send_message(message.chat.id,
                         f"{name}, вы ввели неверное значение\nваш пол:\n1. Мужской\n2. Женский\nВведите цифру с нужным вариантом:")
        bot.register_next_step_handler(message,
                                       lambda msg: gender_teacher(msg, name, teacher_id, teacher_phone_nomber, mail))


def department_teacher(message, name, teacher_id, teacher_phone_nomber, mail, gender):
    department = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute(
        'INSERT INTO teachers (teacher_id, name, phone_number, mail, gender, department) VALUES (?, ?, ?, ?, ?, ?)',
        (teacher_id, name, teacher_phone_nomber, mail, gender, department))
    connection.commit()
    connection.close()
    bot.send_message(message.chat.id, "Вы зарегистрированы!")
    changing_teacher(message, teacher_id)


# ИЗМЕНЕНИЕ ДАННЫХ ДЛЯ СТУДЕНТОВ И ПРЕПОДАВАТЕЛЕЙ
def changing_student(message, student_id):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT name, phone_number, mail, faculty, course, group_number FROM student WHERE student_id = ?",
                   (student_id,))
    info_about_student = cursor.fetchall()
    connection.commit()
    connection.close()
    output = "".join(
        f"Ваши данные:\n1) Имя: {info_about_student[i][0]}\n2) Номер телефона: {info_about_student[i][1]}\n3) Почта: {info_about_student[i][2]}\n4) Факультет: {info_about_student[i][3]}\n5) Курс: {info_about_student[i][4]}\n6) Номер группы: {info_about_student[i][5]}"
        for i in range(len(info_about_student)))
    markup = types.InlineKeyboardMarkup()
    markup.add(types.InlineKeyboardButton("Поменять данные", callback_data="changing_student"))
    bot.send_message(message.chat.id, output, reply_markup=markup)


def nomber_change(message):
    nomber = message.text
    print(nomber)
    student_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT id, faculty FROM discipline")
    info_about_faculty = cursor.fetchall()
    connection.commit()
    connection.close()
    if nomber == "1":
        bot.send_message(message.chat.id, "Пожалуйста, введите свое ФИО")
        bot.register_next_step_handler(message,
                                       lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))
    elif nomber == "2":
        bot.send_message(message.chat.id, "Пожалуйста, введите новый номер телефона:")
        bot.register_next_step_handler(message,
                                       lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))
    elif nomber == "3":
        bot.send_message(message.chat.id, "Пожалуйста, введите новый почтовый адрес:")
        bot.register_next_step_handler(message,
                                       lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))
    elif nomber == "4":
        output = "".join(
            f"{info_about_faculty[i][0]}) {info_about_faculty[i][1]}\n"
            for i in range(len(info_about_faculty)))
        bot.send_message(message.chat.id, f"Укажите ваш Факультет:\n{output} ")
        bot.register_next_step_handler(message,
                                       lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))
    elif nomber == "5":
        bot.send_message(message.chat.id, "Пожалуйста, введите номер курса:")
        bot.register_next_step_handler(message,
                                       lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))
    elif nomber == "6":
        try:
            connection = sqlite3.connect('my_database.db')
            cursor = connection.cursor()
            # Получаем информацию о студенте
            cursor.execute("SELECT faculty, course FROM student WHERE student_id = ?", (student_id,))
            info_about_student = cursor.fetchall()
            if not info_about_student:
                bot.send_message(message.chat.id, "Студент не найден.")
                return
            for i in range(len(info_about_student)):
                faculty = info_about_student[i][0]
                course = info_about_student[i][1]
                # Получаем информацию о группах
                cursor.execute("SELECT id, group_number, faculty, course FROM groups WHERE faculty = ? AND course = ?",
                               (faculty, course))
                info_about_faculty = cursor.fetchall()
            if not info_about_faculty:
                bot.send_message(message.chat.id, "Группы не найдены.")
                return
            output = "".join(
                f"{info_about_faculty[i][0]}) {info_about_faculty[i][1]}, факультет: {info_about_faculty[i][2]}, курс: {info_about_faculty[i][3]}\n"
                for i in range(len(info_about_faculty)))
            bot.send_message(message.chat.id, f"Выберите номер группы:\n{output}")
            bot.register_next_step_handler(message,
                                           lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))

        except Exception as e:
            bot.send_message(message.chat.id, f"Произошла ошибка: {str(e)}")
        finally:
            if 'connection' in locals():
                connection.close()
    else:
        bot.send_message(message.chat.id,
                         "Такого номера нет, попробуйте ещё раз. Введите номер поля, которое хотите изменить:")
        bot.register_next_step_handler(message, nomber_change)



def changing_db_student(message, student_id, nomber, info_about_faculty):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    if nomber == '1':
        new = message.text
        cursor.execute("UPDATE student SET name = ? WHERE student_id= ?", (new, student_id))
        connection.commit()
        connection.close()
        bot.send_message(message.chat.id, "Вы поменяли ФИО")
        changing_student(message, student_id)
    elif nomber == "2":
        new = message.text
        cursor.execute("UPDATE student SET phone_number = ? WHERE student_id= ?", (new, student_id))
        connection.commit()
        connection.close()
        bot.send_message(message.chat.id, "Вы поменяли номер телефона")
        changing_student(message, student_id)
    elif nomber == "3":
        new = message.text
        cursor.execute("UPDATE student SET mail = ? WHERE student_id= ?", (new, student_id))
        connection.commit()
        connection.close()
        bot.send_message(message.chat.id, "Вы поменяли почту")
        changing_student(message, student_id)
    elif nomber == "4":
        try:
            new = int(message.text)
        except ValueError:
            bot.send_message(message.chat.id, "Пожалуйста, введите корректный номер факультета.")
            bot.register_next_step_handler(message,
                                           lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))
            return

        have = False
        for i in range(len(info_about_faculty)):
            if info_about_faculty[i][0] == new:
                have = True
                faculty = info_about_faculty[i][1]
                cursor.execute("UPDATE student SET faculty = ? WHERE student_id= ?", (faculty, student_id))
                connection.commit()
                connection.close()
                bot.send_message(message.chat.id, "Вы поменяли факультет")
                changing_student(message, student_id)

        if not have:
            bot.send_message(message.chat.id, "Вы ввели неверное значение, укажите ваш факультет:")
            bot.register_next_step_handler(message,
                                           lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))

    elif nomber == "5":
        new = message.text
        if new == "1" or new == "2" or new == "3" or new == "4":
            cursor.execute("UPDATE student SET course = ? WHERE student_id= ?", (new, student_id))
            connection.commit()
            connection.close()
            bot.send_message(message.chat.id, "Вы поменяли курс")
            changing_student(message, student_id)
        else:
            bot.send_message(message.chat.id, "Вы ввели неверное значение, укажите ваш курс\nПример: 1")
            bot.register_next_step_handler(message,
                                           lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))
    elif nomber == "6":
        try:
            new = int(message.text)
        except ValueError:
            bot.send_message(message.chat.id, "Вы ввели неверное значение, укажите группу:")
            bot.register_next_step_handler(message,
                                           lambda msg: changing_db_student(msg, student_id, nomber, info_about_faculty))
            return
        have = False
        for i in range(len(info_about_faculty)):
            if info_about_faculty[i][0] == new:
                have = True
                groupp = info_about_faculty[i][1]
                cursor.execute("UPDATE student SET group_number = ? WHERE student_id= ?", (groupp, student_id))
                connection.commit()
                connection.close()
                bot.send_message(message.chat.id, "Вы поменяли номер группы")
                changing_student(message, student_id)
        if not have:
            bot.send_message(message.chat.id, "Вы ввели неверное значение, укажите группу:")
            bot.register_next_step_handler(message,
                                           lambda msg: changing_db_student(msg, student_id, nomber,
                                                                           info_about_faculty))


def changing_teacher(message, teacher_id):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT name, phone_number, mail, department FROM teachers WHERE teacher_id = ?", (teacher_id,))
    info_about_teacher = cursor.fetchall()
    connection.commit()
    connection.close()
    output = "".join(
        f"Ваши данные:\n1) Имя: {info_about_teacher[i][0]}\n2) Номер телефона: {info_about_teacher[i][1]}\n3) Почта: {info_about_teacher[i][2]}\n4) Кафедра: {info_about_teacher[i][3]}"
        for i in range(len(info_about_teacher)))
    markup = types.InlineKeyboardMarkup()
    markup.add(types.InlineKeyboardButton("Поменять данные ", callback_data="changing_teacher"))
    bot.send_message(message.chat.id, output, reply_markup=markup)


def nomber_change_teacher(message):
    nomber = message.text
    teacher_id = message.chat.id
    if nomber == "1":
        bot.send_message(message.chat.id, "Пожалуйста, введите свое ФИО")
        bot.register_next_step_handler(message, lambda msg: changing_db_teacher(msg, teacher_id, nomber))
    elif nomber == "2":
        bot.send_message(message.chat.id, "Пожалуйста, введите новый номер телефона:")
        bot.register_next_step_handler(message, lambda msg: changing_db_teacher(msg, teacher_id, nomber))
    elif nomber == "3":
        bot.send_message(message.chat.id, "Пожалуйста, введите новый почтовый адрес:")
        bot.register_next_step_handler(message, lambda msg: changing_db_teacher(msg, teacher_id, nomber))
    elif nomber == "4":
        bot.send_message(message.chat.id, "Пожалуйста, введите название кафедры:")
        bot.register_next_step_handler(message, lambda msg: changing_db_teacher(msg, teacher_id, nomber))
    else:
        bot.send_message(message.chat.id,
                         "Такого номера нет, попробуйте ещё раз. Введите номер поля, которое хотетите изменить:")
        bot.register_next_step_handler(message, nomber_change_teacher)


def changing_db_teacher(message, teacher_id, nomber):
    new = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    if nomber == '1':
        cursor.execute("UPDATE teachers SET name = ? WHERE teacher_id= ?", (new, teacher_id))
        bot.send_message(message.chat.id, "Вы поменяли ФИО")
        connection.commit()
        changing_teacher(message, teacher_id)
    if nomber == "2":
        cursor.execute("UPDATE teachers SET phone_number = ? WHERE teacher_id= ?", (new, teacher_id))
        bot.send_message(message.chat.id, "Вы поменяли номер телефона")
        connection.commit()
        changing_teacher(message, teacher_id)
    if nomber == "3":
        cursor.execute("UPDATE teachers SET mail = ? WHERE teacher_id= ?", (new, teacher_id))
        bot.send_message(message.chat.id, "Вы поменяли почту")
        connection.commit()
        changing_teacher(message, teacher_id)
    if nomber == "4":
        cursor.execute("UPDATE teachers SET department = ? WHERE teacher_id= ?", (new, teacher_id))
        bot.send_message(message.chat.id, "Вы поменяли кафедру")
        connection.commit()
        changing_teacher(message, teacher_id)
    connection.close()

# Функции для преподавателя
# ДОБАВЛЕНИЕ ПРЕПОДАВАТЕЛЕМ ДАННЫХ В ТАБЛИЦУ ДИСЦИПЛИНА

def add_data_to_table_discipline(message):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    student = cursor.fetchone()[0]
    # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
    if student > 0 and teacher == 0:
        bot.send_message(message.chat.id, f"Вы зарегистрированы как студент. Функция доступна преподавателю.")
    elif student == 0 and teacher > 0:
        bot.send_message(message.chat.id, "Введите название вашей дисциплины:")
        bot.register_next_step_handler(message, lambda msg: to_table_discipline(msg, message.chat.id))
    elif student > 0 and teacher > 0:
        bot.send_message(message.chat.id, f"Вы зарегистрированы и как преподаватель и как студент. Невозможно продолжить работу.")
    elif student == 0 and teacher == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id, "Вы не зарегистрированы. Для того, чтобы продолжить, необходимо пройти регистрацию. ", reply_markup = markup)
    connection.commit()
    connection.close()


def to_table_discipline(message, teacher_id):
    name_discipline = message.text
    bot.send_message(message.chat.id, "Введите название факультета:")
    bot.register_next_step_handler(message, lambda msg: to_table_dis(msg, name_discipline, teacher_id))


def to_table_dis(message, name_discipline, teacher_id):
    name_facyltet = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute('INSERT INTO discipline (name_of_discipline, teacher_id, faculty) VALUES (?, ?, ?)',
                   (name_discipline, teacher_id, name_facyltet))
    connection.commit()
    connection.close()
    bot.send_message(message.chat.id, f"Дисциплина создана")
    select_data_for_teacher(message, teacher_id)


# ВЫВОДИМ СПИСОК ВСЕХ ДИСЦИПЛИН
def select_data_for_teacher(message, teacher_id):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT id, name_of_discipline, faculty FROM discipline WHERE teacher_id = ?", (teacher_id,))
    info_about_discipline = cursor.fetchall()
    connection.commit()
    connection.close()
    if len(info_about_discipline) > 0 :
        output = "".join(
            f"{info_about_discipline[i][0]}) {info_about_discipline[i][1]}, {info_about_discipline[i][2]}\n"
            for i in range(len(info_about_discipline)))
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Поменять наименование дисциплины", callback_data="changing"))
        markup.add(types.InlineKeyboardButton("Добавить дисциплину", callback_data="add"))
        bot.send_message(message.chat.id, f"Дисциплины:\n{output}", reply_markup=markup)
    else:
        bot.send_message(teacher_id, "У вас нет созданных дисциплин.")


# МЕНЯЕМ НАЗВАНИЕ ДИСЦИПЛИНЫ
def nomber_change_discepline(message):
    nomber = message.text
    bot.send_message(message.chat.id, "Пожалуйста, введите новое название дисциплины")
    bot.register_next_step_handler(message, lambda msg: changing_discepline(msg, nomber))


def changing_discepline(message, nomber):
    discepline = message.text
    bot.send_message(message.chat.id, "Пожалуйста, введите новое название факультета:")
    bot.register_next_step_handler(message, lambda msg: changing_disceplineee(msg, nomber, discepline))


def changing_disceplineee(message, nomber, discepline):
    name_facyltet = message.text
    nomber = int(nomber)
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT id, name_of_discipline, faculty FROM discipline WHERE id = ?", (nomber,))
    faculty = cursor.fetchall()
    connection.commit()
    connection.close()
    if not faculty:
        bot.send_message(message.chat.id, "Нет данных с таким идентификатором. Вы ввели неверное число")
        return
    output = "".join(f"{faculty[i][0]}) {faculty[i][1]}\n" for i in range(len(faculty)))
    if output:
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("UPDATE discipline SET name_of_discipline = ?, faculty = ? WHERE id= ?",
                       (discepline, name_facyltet, nomber))
        connection.commit()
        connection.close()
        bot.send_message(message.chat.id, "Вы поменяли данные.")
        select_data_for_teacher(message, message.chat.id)
    else:
        bot.send_message(message.chat.id, "Нет данных с таким идентификатором.")


# ДОБАВЛЯЕМ ГРУППУ

def groap_table(message):
    teacher_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    student = cursor.fetchone()[0]
    # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
    if student > 0 and teacher == 0:
        bot.send_message(message.chat.id, f"Вы зарегистрированы как студент. Функция доступна преподавателю.")
    elif student == 0 and teacher > 0:
        cursor.execute("SELECT DISTINCT faculty FROM discipline WHERE teacher_id = ?", (teacher_id,))
        faculty = cursor.fetchall()
        if not faculty:
            bot.send_message(message.chat.id, "Нет доступных факультетов.")
            return
        output = "".join(f"{i+1}) {faculty[i][0]}\n" for i in range(len(faculty)))
        bot.send_message(message.chat.id,
                         f"Вы создали следующие факультеты:\n{output}\nУкажите номер факультета, в котором создаётся группа:")
        bot.register_next_step_handler(message, lambda msg: to_table_groap(msg, teacher_id, faculty))
    elif student > 0 and teacher > 0:
        bot.send_message(message.chat.id,
                         f"Вы зарегистрированы и как преподаватель и как студент. Невозможно продолжить работу.")
    elif student == 0 and teacher == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "Вы не зарегистрированы. Для того, чтобы продолжить, необходимо пройти регистрацию. ",
                         reply_markup=markup)
    connection.commit()
    connection.close()


def to_table_groap(message, teacher_id, faculty):
    try:
        id = int(message.text)
    except ValueError:
        bot.send_message(message.chat.id, "Пожалуйста, введите корректный номер факультета.")
        bot.register_next_step_handler(message, lambda msg: to_table_groap(msg, teacher_id, faculty))
        return
    have = False
    for i in range(len(faculty)):
        have = True
        facultet = faculty[i][id - 1]
        bot.send_message(message.chat.id, f"Вы выбрали: {facultet}\nВведите номер группы:")
        bot.register_next_step_handler(message, lambda msg, f=facultet: to_groap(msg, teacher_id, f))
    if not have:
        # Если мы вышли из цикла без нахождения факультета
        bot.send_message(message.chat.id, "Неверный номер факультета. Пожалуйста, попробуйте снова.")
        bot.register_next_step_handler(message, lambda msg: to_table_groap(msg, teacher_id, faculty))


def to_groap(message, teacher_id, facultet):
    group = message.text
    bot.send_message(message.chat.id, "Укажите курс:")
    bot.register_next_step_handler(message, lambda msg: to_tableee_groap(msg, teacher_id, facultet, group))


def to_tableee_groap(message, teacher_id, facultet, group):
    cyrs = message.text
    if cyrs == "1" or cyrs == "2" or cyrs == "3" or cyrs == "4":
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute('INSERT INTO groups (group_number, faculty, course) VALUES (?, ?, ?)',
                       (group, facultet, cyrs))
        connection.commit()
        connection.close()
        bot.send_message(message.chat.id, "Группа добавлена, теперь в неё могут добавляться студенты.")
        spisok_grupp(message)
    else:
        bot.send_message(message.chat.id, "Вы ввели неверное значение. Укажите курс:")
        bot.register_next_step_handler(message, lambda msg: to_tableee_groap(msg, teacher_id, facultet, group))


# ВЫВОДИМ СПИСОК ГРУПП
def spisok_grupp(message):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT * FROM groups")
    groups = cursor.fetchall()
    connection.commit()
    connection.close()
    output = "".join(f"{groups[i][0]}) {groups[i][1]}, факультет: {groups[i][2]}, курс: {groups[i][3]}\n" for i in
                     range(len(groups)))
    markup = types.InlineKeyboardMarkup()
    markup.add(types.InlineKeyboardButton("Поменять группу", callback_data="changing_group"))
    markup.add(types.InlineKeyboardButton("Добавить группу", callback_data="add_group"))
    if output:
        bot.send_message(message.chat.id, f"Все группы:\n{output}", reply_markup=markup)
    else:
        bot.send_message(message.chat.id, "Нет данных с таким идентификатором.")


# МЕНЯЕМ НАЗВАНИЕ ГРУППЫ
def nomber_change_group(message):
    nomber = message.text
    bot.send_message(message.chat.id, "Пожалуйста, введите новый номер группы:")
    bot.register_next_step_handler(message, lambda msg: changing_group(msg, nomber))


def changing_group(message, nomber):
    nomber_group = message.text
    bot.send_message(message.chat.id, "Пожалуйста, введите курс:")
    bot.register_next_step_handler(message, lambda msg: changing_grouppp(msg, nomber, nomber_group))


def changing_grouppp(message, nomber, nomber_group):
    cyrs = message.text
    if cyrs == "1" or cyrs == "2" or cyrs == "3" or cyrs == "4":
        nomber = int(nomber)
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("UPDATE groups SET group_number = ?, course = ? WHERE id= ?", (nomber_group, cyrs, nomber))
        connection.commit()
        connection.close()
        bot.send_message(message.chat.id, "вы поменяли данные.")
        spisok_grupp(message)
    else:
        bot.send_message(message.chat.id, "Вы ввели неверное значение. Укажите курс:")
        bot.register_next_step_handler(message, lambda msg: changing_grouppp(msg, nomber, nomber_group))


def create_pie_chart(vals, labels):
    width = 800
    height = 800
    image = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(image)
    total = sum(vals)
    start_angle = 0
    legend_x = 30
    legend_y = 100
    legend_spacing = 30
    font_path = "arial.ttf"
    font = ImageFont.truetype(font_path, 25)
    now = datetime.now()
    months = ["Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
        "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]
    current_month_index = now.month - 1
    current_month = months[current_month_index]
    month_font = ImageFont.truetype(font_path, 40)  # Шрифт для названия месяца
    month_text = f"Месяц: {current_month}"
    bbox = draw.textbbox((0, 0), month_text, font=month_font)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    draw.text(((width - text_width) / 2, 20), month_text, fill="black", font=month_font)

    for i in range(len(vals)):
        end_angle = start_angle + (vals[i] / total) * 360
        color = tuple(int(x) for x in (255 * (i % 2), 100, 100))
        draw.pieslice([100, 200, width - 100, height - 80], start_angle, end_angle, fill=color)
        draw.rectangle([legend_x, legend_y + i * legend_spacing, legend_x + 20, legend_y + i * legend_spacing + 20],
                       fill=color)
        draw.text((legend_x + 24, legend_y + i * legend_spacing),
                  f"{labels[i]}: {vals[i]} ({vals[i] / total * 100:.1f}%)", fill="black", font=font)
        start_angle = end_angle

    return image

#Статистика для препода и студента
@bot.message_handler(commands=['statistics'])
def statis_teacher (message):
    teacher_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (teacher_id,))
    count_teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (teacher_id,))
    count_student = cursor.fetchone()[0]
    if count_teacher > 0 and count_student == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Статистика по отправленной задаче", callback_data="statystic"))
        markup.add(types.InlineKeyboardButton("Общая статистика, средний балл", callback_data="all_statystic"))
        bot.send_message(teacher_id, f"Вас интересуют:", reply_markup=markup)
    elif count_teacher == 0 and count_student > 0:
        cursor.execute("SELECT complete, dont_complete FROM statystic_for_student WHERE student_id = ?", (teacher_id,))
        statystic_for_student = cursor.fetchone()
        if statystic_for_student:
            complete, dont_complete = statystic_for_student
            vals = [complete, dont_complete]
            labels = ["Выполненные задачи", "Не выполненные задачи"]
            pie_chart_image = create_pie_chart(vals, labels)
            buf = BytesIO()
            pie_chart_image.save(buf, format='PNG')
            buf.seek(0)
            # Отправка диаграммы пользователю
            bot.send_photo(teacher_id, photo=buf)
        else:
            bot.send_message(teacher_id, "У вас нет задач.")
    elif count_teacher == 0 and count_student == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(teacher_id, "Чтобы начать работу с ботом необходимо пройти регистрацию. ", reply_markup=markup)
    elif count_teacher > 0 and count_student > 0:
        delete_zapis(message)
    connection.commit()
    connection.close()


#СРЕДНИЙ БАЛЛ СТАТИСТИКА
def all_statystic(message, teacher_id, discipline):
    try:
        discipline_number = int(message.text)
    except ValueError:
        print("Ошибка: Введите корректный номер дисциплины.")
        return
    if discipline_number < 1 or discipline_number > len(discipline):
        print("Ошибка: Номер дисциплины вне диапазона.")
        return
    selected_discipline = discipline[discipline_number - 1][0]
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute(
        "SELECT DISTINCT group_number FROM task_for_student WHERE teacher_id = ? AND name_of_discipline = ?",
        (teacher_id, selected_discipline ))
    group_number = cursor.fetchall()
    connection.commit()
    connection.close()
    if group_number:
        output = "".join(
            f"\n{i + 1}) {group_number[i][0]} "
            for i in
            range(len(group_number)))
        bot.send_message(teacher_id, f"Введите номер группы:\n{output}")
        bot.register_next_step_handler(message, lambda msg: all_statistic_g(msg, teacher_id, selected_discipline, group_number))
    else:
        bot.send_message(teacher_id, "Вы ещё не отправляли задания для этой группы.")

def all_statistic_g(message, teacher_id, selected_discipline, group_number):
    try:
        group = int(message.text)
    except ValueError:
        print("Ошибка: Введите корректный номер группы.")
        return
    if group < 1 or group > len(group_number):
        print("Ошибка: Номер группы вне диапазона.")
        return
    selected_group = group_number[group - 1][0]
    bot.send_message(teacher_id, f"Теперь необходимо задать временной диапазоню\nВведите с какого числа и по какое вывести статистику.\n Пример: 10.10.2024-20.12.2024")
    bot.register_next_step_handler(message,
                                   lambda msg: all_statistic_date(msg, teacher_id, selected_discipline, selected_group))

def all_statistic_date(message, teacher_id, selected_discipline, selected_group):
    date = message.text.strip()
    try:
        date1_str, date2_str = date.split('-')
        date1_str = date1_str.strip()
        date2_str = date2_str.strip()

        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute(
            """
            SELECT student_id, name_student, task_id, complete 
            FROM task_list 
            WHERE teacher_id = ? 
              AND name_of_discipline = ? 
              AND group_number = ? 
              AND send_teacher_for_student_date >= ? 
              AND send_teacher_for_student_date <= ?
            """,
            (teacher_id, selected_discipline, selected_group, date1_str, date2_str))
        all_results = cursor.fetchall()
        student_stats = {}
        for record in all_results:
            student_id, name_student, task_id, complete = record
            if name_student not in student_stats:
                student_stats[name_student] = {
                    'completed_tasks': [],
                    'incompleted_tasks': [],
                    'marks': []
                }
            if complete == 1:
                # Получаем оценку для выполненной задачи
                cursor.execute(
                    """
                    SELECT mark 
                    FROM teacher_comment 
                    WHERE teacher_id = ? 
                      AND name_of_discipline = ? 
                      AND group_number = ? 
                      AND task_id = ? 
                      AND student_id = ?
                    """,
                    (teacher_id, selected_discipline, selected_group, task_id, student_id)
                )
                mark_result = cursor.fetchone()
                mark = mark_result[0] if mark_result else None

                if mark is not None:
                    student_stats[name_student]['completed_tasks'].append((task_id, mark))
                    student_stats[name_student]['marks'].append(mark)
            else:
                student_stats[name_student]['incompleted_tasks'].append(task_id)
        connection.commit()
        connection.close()
        doc = Document()
        doc.add_heading(f'{selected_discipline}\nГруппа {selected_group}\n{date1_str} - {date2_str}', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Имя студента'
        hdr_cells[1].text = 'Выполненные задачи с оценками'
        hdr_cells[2].text = 'Невыполненные задачи'
        hdr_cells[3].text = 'Средний балл'

        for name_student, stats in student_stats.items():
            completed_tasks_str = ', '.join([f'Задача {task_id}: {mark}' for task_id, mark in stats['completed_tasks']])
            incompleted_tasks_str = ', '.join([f'Задача {task_id}' for task_id in stats['incompleted_tasks']])
            average_mark = sum(stats['marks']) / (len(stats['marks']) + len(stats['incompleted_tasks'])) if stats['marks'] else 0
            row_cells = table.add_row().cells

            row_cells[0].text = str(name_student)
            row_cells[1].text = completed_tasks_str
            row_cells[2].text = incompleted_tasks_str
            row_cells[3].text = f'{average_mark:.2f}'
        document_name = f'statistics.docx'
        doc.save(document_name)
        bot.send_document(teacher_id, open(document_name, "rb"))
        os.remove(document_name)
    except sqlite3.Error as e:
        bot.send_message(teacher_id, (f"Ошибка при работе с базой данных: {e}"))
    except ValueError:
        bot.send_message(teacher_id, "Ошибка: Неверный формат даты. Пожалуйста, убедитесь, что даты введены корректно.")


# ВЫВОДИМ СПИСОК СТУДЕНТОВ СДАВШИХ И НЕ СДАВШИХ РАБОТУ
def statystics(message, teacher_id):
    try:
        nomber = int(message.text)
    except ValueError:
        bot.send_message(teacher_id,
                         "Неверный номер. Попробуйте ещё раз:")
        bot.register_next_step_handler(message,
                                       lambda msg: statystics(msg, teacher_id))
        return
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute(
        "SELECT name_student, group_number, the_task_for_student, name_of_discipline FROM task_list WHERE teacher_id = ? AND task_id = ? AND complete IS NULL",
        (teacher_id, nomber))
    info_dont_complete_task = cursor.fetchall()
    cursor.execute(
        "SELECT id, name_of_discipline, group_number, the_task_for_student, send_time, send_date FROM task_for_student WHERE teacher_id = ? AND statys = 1 ",
        (teacher_id,))
    info_send_task = cursor.fetchall()
    cursor.execute(
        "SELECT name_student, group_number, task_time, date, the_task_for_student, name_of_discipline FROM task_list WHERE teacher_id = ? AND task_id = ? AND complete = 1 ",
        (teacher_id, nomber))
    info_complete_task = cursor.fetchall()
    cursor.execute("SELECT COUNT (*) FROM task_list WHERE teacher_id = ? AND task_id = ? AND complete IS NULL",
                   (teacher_id, nomber))
    count_dont_complete_task = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM task_list WHERE teacher_id = ? AND task_id = ? AND complete = 1 ",
                   (teacher_id, nomber))
    count_complete_task = cursor.fetchone()[0]
    connection.commit()
    connection.close()
    have = False
    for i in range(len(info_send_task)):
        if info_send_task[i][0] == nomber:
            have = True
            if info_complete_task:
                output = "".join(
                    f"\n{info_complete_task[i][0]}- {info_complete_task[i][1]}, {info_complete_task[i][5]}\nотправил(-ла) решение по задаче:\n{info_complete_task[i][4]}\nВ {info_complete_task[i][2]} {info_complete_task[i][3]}"
                    for i in
                    range(len(info_complete_task)))
                bot.send_message(message.chat.id, f"РЕШЕНИЕ ОТПРАВИЛИ {count_complete_task} студента(-ов):{output}")
                if info_dont_complete_task:
                    out = "".join(
                        f"\n{info_dont_complete_task[i][0]}- {info_dont_complete_task[i][1]}"
                        for i in
                        range(len(info_dont_complete_task)))
                    bot.send_message(message.chat.id, f"НЕ ОТПРАВИЛИ {count_dont_complete_task} студента(-ов): {out}")
            else:
                bot.send_message(message.chat.id, f"Нет решённых задач от студентов.")
    if not have:
        bot.send_message(teacher_id, "Такого номера нет.")

#ТПРАВИТЬ УВЕДОМЛЕНИЕ СТУДЕНТАМ
def send_message_for_student(message):
    teacher_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    student = cursor.fetchone()[0]
    # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
    if student > 0 and teacher == 0:
        bot.send_message(message.chat.id, f"Вы зарегистрированы как студент. Функция доступна преподавателю.")
    elif student == 0 and teacher > 0:
        cursor.execute("SELECT DISTINCT faculty FROM discipline WHERE teacher_id = ?", (teacher_id,))
        count_teacher = cursor.fetchall()
        for i in count_teacher:
            faculty = i[0]
            cursor.execute("SELECT DISTINCT group_number FROM student WHERE faculty = ?", (faculty,))
            info_ab = cursor.fetchall()
            output = "".join(
                f"{i + 1}) {info_ab[i][0]} " for i in
                range(len(info_ab)))
            bot.send_message(teacher_id, f"{output}\nВведите номер группы:")
            bot.register_next_step_handler(message,
                                       lambda msg: send_message_for_studenttt(msg, teacher_id, info_ab))
        if not count_teacher:
            bot.send_message(teacher_id, "Нет студентов, которые зарегистрированы в вашем факультете")
    elif student > 0 and teacher > 0:
        bot.send_message(message.chat.id,
                         f"Вы зарегистрированы и как преподаватель и как студент. Невозможно продолжить работу.")
    elif student == 0 and teacher == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "Вы не зарегистрированы. Для того, чтобы продолжить, необходимо пройти регистрацию. ",
                         reply_markup=markup)
    connection.commit()
    connection.close()

def send_message_for_studenttt(message, teacher_id, info_ab):
    try:
        nomber = int(message.text)
    except ValueError:
        bot.send_message(teacher_id,
                         "Неверный номер. Попробуйте ещё раз:")
        bot.register_next_step_handler(message,
                                       lambda msg: send_comment(msg, teacher_id))
        return

    if nomber < 1 or nomber > len(info_ab):
        bot.send_message(message.chat.id,
                         f'Введённый номер задачи вне диапазона.')
        return
    groupp = info_ab[nomber - 1][0]
    bot.send_message(teacher_id, f"Введите сообщение, которое хотите отправить студентам группы {groupp} :")
    bot.register_next_step_handler(message,
                                       lambda msg: send_message_for_studentttt(msg, teacher_id, groupp))

def send_message_for_studentttt(message, teacher_id, groupp):
    message_for_student = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute(""" SELECT student_id 
                        FROM student 
                        WHERE group_number = ? """, (groupp,))
    for_student = cursor.fetchall()
    if len(for_student) > 0:
        cursor.execute(""" SELECT name, department 
                                FROM teachers 
                                WHERE teacher_id = ? """, (teacher_id,))
        info_ab_teacher = cursor.fetchone()
        if info_ab_teacher:
            name, department = info_ab_teacher
            for student in for_student:
                student_id = student[0]
                send_message_ga(student_id, f"Сообщение от преподавателя\n{name}\n{department}:\n{message_for_student}")
            bot.send_message(teacher_id, "Отправлено")
        else:
            bot.send_message(teacher_id, "Информация о преподавателе не найдена.")
    else:
        bot.send_message(teacher_id, f"В группе {groupp} нет зарегистрированных студентов.")
    connection.close()

def complete_task(message):
    teacher_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    student = cursor.fetchone()[0]
    # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
    if student > 0 and teacher == 0:
        bot.send_message(message.chat.id, f"Вы зарегистрированы как студент. Функция доступна преподавателю.")
    elif student == 0 and teacher > 0:
        cursor.execute(
            "SELECT id, name_of_discipline, group_number, the_task_for_student, send_time, send_date FROM task_for_student WHERE teacher_id = ? AND statys = 1",
            (teacher_id,))
        info_send_task = cursor.fetchall()
        if info_send_task:
            output = "".join(
                f"\n{info_send_task[i][0]}) {info_send_task[i][2]} {info_send_task[i][1]}\nЗадача:\n{info_send_task[i][3]}\nотправлено в {info_send_task[i][4]} {info_send_task[i][5]}"
                for i in
                range(len(info_send_task)))
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Отправить оценку", callback_data="send_mark"))
            bot.send_message(message.chat.id, f"Вы отправили задачи:{output}", reply_markup=markup)
        else:
            bot.send_message(message.chat.id, f"У вас нет отправленных задач.")
    elif student > 0 and teacher > 0:
        bot.send_message(message.chat.id,
                         f"Вы зарегистрированы и как преподаватель и как студент. Невозможно продолжить работу.")
    elif student == 0 and teacher == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "Вы не зарегистрированы. Для того, чтобы продолжить, необходимо пройти регистрацию. ",
                         reply_markup=markup)
    connection.commit()
    connection.close()

def send_comment(message, teacher_id):
    try:
        nomber = int(message.text)
    except ValueError:
        bot.send_message(teacher_id,
                         "Неверный номер. Попробуйте ещё раз:")
        bot.register_next_step_handler(message,
                                       lambda msg: send_comment(msg, teacher_id))
        return
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute(
        "SELECT id, name_of_discipline, group_number, the_task_for_student, send_time, send_date FROM task_for_student WHERE teacher_id = ? AND statys = 1 ",
        (teacher_id,))
    info_send_task = cursor.fetchall()
    have = False
    for i in range(len(info_send_task)):
        if info_send_task[i][0] == nomber:
            have = True
            cursor.execute(
                "SELECT id, name_student, group_number, task_time, date, the_task_for_student, name_of_discipline FROM task_list WHERE teacher_id = ? AND task_id = ? AND complete = 1 AND send_mark_date IS NULL",
                (teacher_id, nomber))
            info_complete_task = cursor.fetchall()
            if info_complete_task:
                output = "".join(
                    f"\n{info_complete_task[i][0]}) {info_complete_task[i][1]}- {info_complete_task[i][2]}, {info_complete_task[i][6]}\nотправил(-ла) решение по задаче:\n{info_complete_task[i][5]}\nВ {info_complete_task[i][3]} {info_complete_task[i][4]}"
                    for i in
                    range(len(info_complete_task)))
                bot.send_message(message.chat.id, f"РЕШЕНИЕ ОТПРАВИЛИ:\n{output}")
                bot.send_message(message.chat.id, f"Введите номер решения, по которому хотите отправить оценку:")
                bot.register_next_step_handler(message,
                                               lambda msg: send_mark(msg, teacher_id, nomber, info_complete_task))
            else:
                bot.send_message(message.chat.id, f"Нет решённых задач от студентов.")
    connection.commit()
    connection.close()
    if not have:
        bot.send_message(teacher_id, "Такого номера нет.")


def send_mark(message, teacher_id, nomber, info_complete_task):
    try:
        id = int(message.text)
    except ValueError:
        bot.send_message(teacher_id,
                         "Неверный номер. Попробуйте ещё раз:")
        bot.register_next_step_handler(message,
                                       lambda msg: send_mark(msg, teacher_id, nomber, info_complete_task))
        return
    have = False
    for i in range(len(info_complete_task)):
        if info_complete_task[i][0] == id:
            have = True
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("2 - неудовлетворительно", callback_data=f"neyd{id}-{nomber}"))
            markup.add(
                types.InlineKeyboardButton("3 - удовлетворительно", callback_data=f"ydovletvoritelno{id}-{nomber}"))
            markup.add(types.InlineKeyboardButton("4 - хорошо", callback_data=f"horosho{id}-{nomber}"))
            markup.add(types.InlineKeyboardButton("5 - отлично", callback_data=f"otlichno{id}-{nomber}"))
            bot.send_message(message.chat.id, f"Введите оценку:", reply_markup=markup)
    if not have:
        bot.send_message(teacher_id, "Такого номера нет.")


def ocenka(message, teacher_id, mark, id, nomber):
    comment = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute(
        "SELECT task_id, student_id, group_number, the_task_for_student, name_of_discipline FROM task_list WHERE teacher_id = ? AND id = ? AND complete = 1",
        (teacher_id, id))
    complete = cursor.fetchall()
    now = datetime.now()
    current_date = now.strftime("%d.%m")
    current_time = now.strftime("%H:%M")
    for task in complete:
        task_id, student_id, group_number, the_task_for_student, name_of_discipline = task
        cursor.execute(
            'INSERT INTO teacher_comment (task_id, student_id, teacher_id, name_of_discipline, the_task_for_student, send_time, date, group_number, comment, mark) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
            (task_id, student_id, teacher_id, name_of_discipline, the_task_for_student, current_time, current_date,
             group_number, comment, mark))
        connection.commit()
        connection.close()
        bot.send_message(teacher_id,
                         f"Ваш комментарий:\nОценка {mark}\n{comment}\nБудет отправлен студенту. Вам придёт уведомление")
        send_coment_teacher(current_date,current_time)

#Получаем информацию о студентах

def info(message):
    teacher_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    student = cursor.fetchone()[0]
    # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
    if student > 0 and teacher == 0:
        bot.send_message(message.chat.id, f"Вы зарегистрированы как студент. Функция доступна преподавателю.")
    elif student == 0 and teacher > 0:
        cursor.execute("SELECT DISTINCT faculty FROM discipline WHERE teacher_id = ?", (teacher_id,))
        count_teacher = cursor.fetchall()
        for i in count_teacher:
            faculty = i[0]
            cursor.execute("SELECT DISTINCT name, phone_number, mail, group_number FROM student WHERE faculty = ?", (faculty,))
            info_ab = cursor.fetchall()
            output = "".join(
                f"{i + 1}) {info_ab[i][0]}, {info_ab[i][1]}, {info_ab[i][2]} ГРУППА {info_ab[i][3]} " for i in
                range(len(info_ab)))
            bot.send_message(teacher_id, f"{output}")
        if not count_teacher:
            bot.send_message(teacher_id, "Нет студентов, которые зарегистрированы в вашем факультете")
    elif student > 0 and teacher > 0:
        bot.send_message(message.chat.id,
                         f"Вы зарегистрированы и как преподаватель и как студент. Невозможно продолжить работу.")
    elif student == 0 and teacher == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "Вы не зарегистрированы. Для того, чтобы продолжить, необходимо пройти регистрацию. ",
                         reply_markup=markup)
    connection.commit()
    connection.close()

# Удаляем учётную запись

def delete_zapis(message):
    user_id = message.chat.id
    # ПРОВЕРКА РЕГИСТРАЦИИ ПОЛЬЗОВАТЕЛЯ
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (user_id,))
    count_teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (user_id,))
    count_student = cursor.fetchone()[0]
    if count_teacher > 0 and count_student > 0:
        bot.send_message(message.chat.id,
                         f"У вас две учётные записи. Вы зарегистрированы и как студент и как преподаватель.\nКакую хотите удалить?\n1. Преподаватель\n2. Студент ")
        bot.register_next_step_handler(message, lambda msg: teacher_or_student_account(msg, user_id))
    elif count_teacher > 0 and count_student == 0:
        cursor.execute("SELECT name, phone_number, mail, department FROM teachers WHERE teacher_id = ?", (user_id,))
        info_about_teacher = cursor.fetchall()
        connection.commit()
        connection.close()
        output = "".join(
            f"Ваши данные:\n1) Имя: {info_about_teacher[i][0]}\n2) Номер телефона: {info_about_teacher[i][1]}\n3) Почта: {info_about_teacher[i][2]}\n4) Кафедра: {info_about_teacher[i][3]}"
            for i in range(len(info_about_teacher)))
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Да", callback_data="delete"))
        markup.add(types.InlineKeyboardButton("Нет", callback_data="ne_delete"))
        bot.send_message(message.chat.id, f"У вас есть 1 учётная запись,\n{output}\n хотите удалить?",
                         reply_markup=markup)
    elif count_teacher == 0 and count_student > 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Да", callback_data="deletee"))
        markup.add(types.InlineKeyboardButton("Нет", callback_data="ne_delete"))
        cursor.execute(
            "SELECT name, phone_number, mail, faculty, course, group_number FROM student WHERE student_id = ?",
            (user_id,))
        info_about_student = cursor.fetchall()
        connection.commit()
        connection.close()
        output = "".join(
            f"Ваши данные:\n1) Имя: {info_about_student[i][0]}\n2) Номер телефона: {info_about_student[i][1]}\n3) Почта: {info_about_student[i][2]}\n4) Факультет: {info_about_student[i][3]}\n5) Курс: {info_about_student[i][4]}\n6) Номер группы: {info_about_student[i][5]}"
            for i in range(len(info_about_student)))
        bot.send_message(message.chat.id, f"У вас есть 1 учётная запись,\n{output}\n хотите удалить?",
                         reply_markup=markup)
    else:
        connection.commit()
        connection.close()
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "У вас нет учётной записи. Хотите зарегистрироваться?",
                         reply_markup=markup)


def teacher_or_student_account(message, user_id):
    statys = message.text
    if statys == "1":
        statys = int(statys)
        delete_user(message, user_id, statys)
    elif statys == "2":
        statys = int(statys)
        delete_user(message, user_id, statys)
    else:
        bot.send_message(message.chat.id, "Вы ввели неверное значение")


def delete_user(message, user_id, statys):
    if statys == 1:
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("DELETE FROM teachers WHERE teacher_id = ?", (user_id,))
        connection.commit()
        connection.close()
        bot.send_message(message.chat.id,
                         f"Ваша учётная запись удалена")
    if statys == 2:
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("DELETE FROM student WHERE student_id= ?", (user_id,))
        connection.commit()
        connection.close()
        bot.send_message(message.chat.id,
                         f"Ваша учётная запись удалена")


# ДОБАВЛЕНИЕ РЕГУЛЯРНОЙ (СТУДЕНТ) И НЕРЕГЕГУЛЯРНОЙ ЗАДАЧИ (СТУДЕНТ И ПРЕПОДАВАТЕЛЬ)

def new_task(message, text):
    regular = None
    user_id = message.chat.id
    # ПРОВЕРКА РЕГИСТРАЦИИ ПОЛЬЗОВАТЕЛЯ
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (user_id,))
    count_teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (user_id,))
    count_student = cursor.fetchone()[0]
    connection.commit()
    connection.close()
    if count_teacher > 0 and count_student > 0:
        delete_zapis(message)
    elif count_teacher > 0 and count_student == 0:
        if text == "add_regular_task":
            bot.send_message(message.chat.id, "Функция доступна только студентам.")
        else:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Сейчас", callback_data="now"))
            markup.add(types.InlineKeyboardButton("Ко времени", callback_data="time"))
            bot.send_message(message.chat.id,
                             "Когда отправить задачу?",
                             reply_markup=markup)
    elif count_teacher == 0 and count_student > 0:
        statys = 2
        print(message.text)
        if text == "add_regular_task":
            send = "time"
            regular = True
            bot.send_message(message.chat.id, "Какую задачу хотите запланировать? Введите название:")
            bot.register_next_step_handler(message, lambda msg: whattime(msg, message.chat.id, regular, statys, send))
        else:
            regular = False
            send = "time"
            bot.send_message(message.chat.id, "Какую задачу хотите запланировать? Введите название:")
            bot.register_next_step_handler(message, lambda msg: whattime(msg, message.chat.id, regular, statys, send))
    else:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "У вас нет учётной записи. Для того, чтобы создать задачу необходимо зарегистрироваться.",
                         reply_markup=markup)

def whattime(message, user_id, regular, statys, send):
    regular = regular
    task_plan = message.text
    if send == 'time':
    # ПРЕПОДАВАТЕЛЬ
        if statys == 1:
            bot.send_message(message.chat.id, "В какое время отправить задачу? (Например - 13:30)")
            bot.register_next_step_handler(message, lambda msg: save_time(msg, task_plan, user_id, regular, statys, send))
        # СТУДЕНТ
        elif statys == 2:
            bot.send_message(message.chat.id, "На какое время? (Например - 13:30)")
            bot.register_next_step_handler(message, lambda msg: save_time(msg, task_plan, user_id, regular, statys, send))
    if send == "now":
        what_time = 0
        save_task(message, task_plan, user_id, what_time, regular, statys, send)



def save_time(message, task_plan, user_id, regular, statys, send):
    regular = regular
    what_time = message.text
    # Проверяем, что часы в диапазоне от 0 до 23 и минуты от 0 до 59
    try:
        hours, minutes = map(int, what_time.split(':'))
        if 0 <= hours < 24 and 0 <= minutes < 60:
            bot.send_message(message.chat.id,
                             "На какую дату хотите запланировать? Пожалуйста, используйте формат: ДД.ММ \n(Например - 12.07)")
            bot.register_next_step_handler(message,
                                           lambda msg: save_task(msg, task_plan, user_id, what_time, regular, statys, send))
        else:
            bot.send_message(message.chat.id,
                             f"Неверный формат. Пожалуйста, используйте формат: ЧЧ:ММ \n Пример (13:30)")
            bot.register_next_step_handler(message, lambda msg: save_time(msg, task_plan, user_id, regular, statys, send))

    except ValueError:
        bot.send_message(user_id,
                         "Неверный формат. Пожалуйста, используйте формат: ЧЧ:ММ \n Пример (13:30)")
        bot.register_next_step_handler(message, lambda msg: save_time(msg, task_plan, user_id, regular, statys, send))


def save_task(message, task_plan, user_id, what_time, regular, statys, send):
    if send == "time":
        date_time = message.text
        if statys == 1:
            try:
                days, month = map(int, date_time.split('.'))
                print(month, days)
                if 1 <= month <= 12 and 1 <= days < 32:
                    connection = sqlite3.connect('my_database.db')
                    cursor = connection.cursor()
                    cursor.execute("SELECT id, name_of_discipline,faculty FROM discipline WHERE teacher_id = ?", (user_id,))
                    discipline = cursor.fetchall()
                    connection.commit()
                    connection.close()
                    output = "".join(f"{discipline[i][0]}) {discipline[i][1]}, факультет: {discipline[i][2]}\n" for i in
                                     range(len(discipline)))
                    if output:
                        bot.send_message(message.chat.id, f"Выберите дисциплину:\n{output}")
                        bot.register_next_step_handler(message,
                                                       lambda msg: discipline_number_statys_teacher_1(msg, task_plan,
                                                                                                      user_id, what_time,
                                                                                                      date_time,
                                                                                                      discipline))
                    else:
                        bot.send_message(message.chat.id, f"Вы ещё не добавили дисциплину.")
                else:
                    bot.send_message(message.chat.id,
                                     "Неверный формат. Пожалуйста, используйте формат: ДД.ММ \n(Например - 12.07)")
                    bot.register_next_step_handler(message,
                                                   lambda msg: save_task(msg, task_plan, user_id, what_time, regular,
                                                                         statys, send))
            except ValueError:
                bot.send_message(user_id,
                                 "Неверный формат. Пожалуйста, используйте формат: ДД.ММ \n(Например - 12.07)")
                bot.register_next_step_handler(message,
                                               lambda msg: save_task(msg, task_plan, user_id, what_time, regular, statys, send))
            except Exception as e:
                bot.send_message(user_id, f"Произошла ошибка: {str(e)}")
        if statys == 2:
            try:
                days, month = map(int, date_time.split('.'))
                if 1 <= month <= 12 and 1 <= days < 32:
                    count_regular_task = 0
                    connection = sqlite3.connect('my_database.db')
                    cursor = connection.cursor()
                    cursor.execute(
                        'INSERT INTO tasks (user_id, task, task_time, date, regular_task, count_regular_task) VALUES (?, ?, ?, ?, ?, ?)',
                        (user_id, task_plan, what_time, date_time, regular, count_regular_task))
                    cursor.execute("SELECT COUNT (*) FROM statystic_for_student WHERE student_id = ?", (user_id,))
                    count_statystic_for_student = cursor.fetchone()[0]
                    if count_statystic_for_student == 0:
                        complete = 0
                        dont_complete = 0
                        all_tasks = 0
                        cursor.execute('INSERT INTO statystic_for_student (student_id, complete, dont_complete, all_tasks) VALUES (?, ?, ?, ?)',
                            (user_id, complete, dont_complete, all_tasks))
                    connection.commit()
                    connection.close()
                    bot.send_message(message.chat.id, "Задача добавлена!")
                else:
                    bot.send_message(message.chat.id,
                                     "Неверный формат. Пожалуйста, используйте формат: ДД.ММ \n(Например - 12.07)")
                    bot.register_next_step_handler(message,
                                                   lambda msg: save_task(msg, task_plan, user_id, what_time, regular,
                                                                         statys, send))
            except ValueError:
                bot.send_message(user_id,
                                 "Неверный формат. Пожалуйста, используйте формат: ДД.ММ \n(Например - 12.07)")
                bot.register_next_step_handler(message,
                                               lambda msg: save_task(msg, task_plan, user_id, what_time, regular, statys, send))
            except Exception as e:
                bot.send_message(user_id, f"Произошла ошибка: {str(e)}")
    if send == "now":
        date_time = 0
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("SELECT id, name_of_discipline,faculty FROM discipline WHERE teacher_id = ?", (user_id,))
        discipline = cursor.fetchall()
        connection.commit()
        connection.close()
        output = "".join(f"{discipline[i][0]}) {discipline[i][1]}, факультет: {discipline[i][2]}\n" for i in
                         range(len(discipline)))
        if output:
            bot.send_message(message.chat.id, f"Выберите дисциплину:\n{output}")
            bot.register_next_step_handler(message,
                                           lambda msg: discipline_number_statys_teacher_1(msg, task_plan,
                                                                                          user_id, what_time,
                                                                                          date_time,
                                                                                          discipline))
        else:
            bot.send_message(message.chat.id, f"Вы ещё не добавили дисциплину.")

# ДАЛЬШЕ ТОЛЬКО ДЛЯ ПРЕПОДАВАТЕЛЯ
def discipline_number_statys_teacher_1(message, task_plan, user_id, what_time, date_time, discipline):
    try:
        nomer = int(message.text)
    except ValueError:
        bot.send_message(user_id,
                         "Неверный номер. Попробуйте ещё раз:")
        bot.register_next_step_handler(message,
                                       lambda msg: discipline_number_statys_teacher_1(msg, task_plan, user_id,
                                                                                      what_time, date_time, discipline))
        return
    have = False
    for i in range(len(discipline)):
        if discipline[i][0] == nomer:
            have = True
            facultet = discipline[i][2]
            name_of_discipline = discipline[i][1]
            connection = sqlite3.connect('my_database.db')
            cursor = connection.cursor()
            cursor.execute("SELECT id, group_number, faculty, course FROM groups WHERE faculty = ?", (facultet,))
            group = cursor.fetchall()
            connection.commit()
            connection.close()
            output = "".join(f"{group[i][0]}) {group[i][1]}, факультет: {group[i][2]}, {group[i][3]}\n" for i in
                             range(len(group)))
            if output:
                bot.send_message(message.chat.id, f"Выберите группу:\n{output}")
                bot.register_next_step_handler(message,
                                               lambda msg: group_number_statys_teacher_1(msg, task_plan, user_id,
                                                                                         what_time, date_time,
                                                                                         name_of_discipline, facultet,
                                                                                         group))
            else:
                bot.send_message(message.chat.id, f"Вы ещё не добавили группу.")
    if not have:
        bot.send_message(user_id,
                         "Неверный номер. Попробуйте ещё раз:")
        bot.register_next_step_handler(message,
                                       lambda msg: discipline_number_statys_teacher_1(msg, task_plan, user_id,
                                                                                      what_time, date_time, discipline))


def group_number_statys_teacher_1(message, task_plan, user_id, what_time, date_time, name_of_discipline, facultet,
                                  group):
    try:
        id = int(message.text)
    except ValueError:
        bot.send_message(message.chat.id, "Пожалуйста, введите корректный номер группы.")
        bot.register_next_step_handler(message,
                                       lambda msg: group_number_statys_teacher_1(msg, task_plan, user_id, what_time,
                                                                                 date_time, name_of_discipline,
                                                                                 facultet, group))
        return
    have = False
    for i in range(len(group)):
        if group[i][0] == id:
            have = True
            group_number = group[i][1]
            course = group[i][3]
            bot.send_message(message.chat.id, "Загрузите документ с заданием")
            bot.register_next_step_handler(message,
                                           lambda msg: document_number_statys_teacher_1(msg, task_plan, user_id,
                                                                                        what_time,
                                                                                        date_time,
                                                                                        name_of_discipline, facultet,
                                                                                        group_number, course))
    if not have:
        bot.send_message(message.chat.id, "Неверный номер. Попробуйте ещё раз:")
        bot.register_next_step_handler(message,
                                       lambda msg: group_number_statys_teacher_1(msg, task_plan, user_id, what_time,
                                                                                 date_time, name_of_discipline,
                                                                                 facultet, group))


def document_number_statys_teacher_1(message, task_plan, user_id, what_time, date_time, name_of_discipline, facultet,
                                     group_number, course):
    if message.document:
        file_name = message.document.file_name
        file_info = bot.get_file(message.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        with open(file_name, 'wb') as new_file:
            new_file.write(downloaded_file)
        if what_time == 0 and date_time == 0:
            statys = 1
            now = datetime.now()
            new_time = now - timedelta(minutes=1)
            date_time = now.strftime("%d.%m")
            what_time = new_time.strftime("%H:%M")
            connection = sqlite3.connect('my_database.db')
            cursor = connection.cursor()
            cursor.execute(
                'INSERT INTO task_for_student (send_date, send_time, name_of_discipline, the_task_for_student, document, group_number, teacher_id, faculty, course, statys) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)',
                (date_time, what_time, name_of_discipline, task_plan, file_name, group_number, user_id, facultet,
                 course, statys))
            connection.commit()
            connection.close()
            bot.send_message(message.chat.id,
                             "Вся информация загружена.\nПосле отправки задания придёт уведомление.")
            # ВЫВОД ЗАГРУЖЕНОЙ ИНФОРМАЦИИ
            conn = sqlite3.connect('my_database.db')
            cursor = conn.cursor()
            cursor.execute("""
                        SELECT student_id 
                        FROM student 
                        WHERE group_number = ? """, (group_number,))
            for_student = cursor.fetchall()
            for student in for_student:
                student_id = student[0]
                send_message_ga(student_id, f"{name_of_discipline}\nЗАДАНИЕ:\n{task_plan}\n")
                bot.send_document(student_id, open(f"{file_name}", "rb"))
                cursor.execute("""
                                            SELECT name 
                                            FROM student 
                                            WHERE student_id = ? """, (student_id,))
                name_student = cursor.fetchall()
                for names in name_student:
                    student_name = names[0]
                    cursor.execute("""
                                                SELECT id 
                                                FROM task_for_student 
                                                WHERE send_date = ? AND send_time = ? AND name_of_discipline = ? AND the_task_for_student = ? AND document = ? AND group_number = ? AND teacher_id = ? AND faculty = ? AND course = ? AND statys = ?""", (date_time, what_time, name_of_discipline, task_plan, file_name, group_number, user_id, facultet,
                 course, statys))
                    for_id = cursor.fetchall()
                    for ids in for_id:
                        id = ids[0]
                        cursor.execute(
                            'INSERT INTO task_list (name_student, task_id, student_id, teacher_id, name_of_discipline, the_task_for_student, group_number, send_teacher_for_student_date) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                            (student_name, id, student_id, user_id, name_of_discipline, task_plan,
                             group_number, date_time))
            send_message_ga(user_id,
                            f"{name_of_discipline}\nЗадача: {task_plan}\nотправлена студентам группы {group_number}")
            os.remove(file_name)
        else:
            conn = sqlite3.connect('my_database.db')
            cursor = conn.cursor()
            cursor.execute(
                'INSERT INTO task_for_student (send_date, send_time, name_of_discipline, the_task_for_student, document, group_number, teacher_id, faculty, course) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)',
                (date_time, what_time, name_of_discipline, task_plan, file_name, group_number, user_id, facultet,
                 course))
            bot.send_message(message.chat.id, "Вся информация загружена.\nПосле отправки задания придёт уведомление.")
            # ВЫВОД ЗАГРУЖЕНОЙ ИНФОРМАЦИИ
        conn.commit()
        conn.close()
    else:
        bot.send_message(message.chat.id, "Пожалуйста, загрузите документ.")
        bot.register_next_step_handler(message,
                                       lambda msg: document_number_statys_teacher_1(msg, task_plan, user_id, what_time,
                                                                                    date_time, name_of_discipline,
                                                                                    facultet, group_number, course))


# ДЛЯ СТУДЕНТОВ ВСЕ ЗАДАЧИ ОТ ПРЕПОДАВАТЕЛЯ/ ОТПРАВИТЬ РЕШЕНИЕ

def task_list(message):
    student_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    count_teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    count_student = cursor.fetchone()[0]
    if count_teacher > 0 and count_student > 0:
        delete_zapis(message)
    elif count_teacher > 0 and count_student == 0:
        bot.send_message(message.chat.id, "Функция доступна только студентам.")
    elif count_teacher == 0 and count_student > 0:
        cursor.execute(
            "SELECT id, name_of_discipline, the_task_for_student FROM task_list WHERE student_id = ? AND complete IS NULL ",
            (student_id,))
        tasks = cursor.fetchall()
        output = "".join(f"{i + 1}) {tasks[i][1]}\nЗАДАЧА:\n{tasks[i][2]}, не выполнено\n" for i in range(len(tasks)))
        if tasks:
            markup = types.InlineKeyboardMarkup()
            markup.add(
                types.InlineKeyboardButton("Отправить решение преподавателю", callback_data="send_completed_task"))
            bot.send_message(message.chat.id, f"Все ваши задачи:\n{output}",
                             reply_markup=markup)
        else:
            bot.send_message(message.chat.id, f'У вас нет задач от преподавателя.')
    else:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "У вас нет учётной записи. Для того, чтобы создать задачу необходимо зарегистрироваться.",
                         reply_markup=markup)
    connection.commit()
    connection.close()

def send_task_for_teacher(message, student_id):
    try:
        nomber = int(message.text)
    except ValueError:
        bot.send_message(student_id,
                         "Неверный номер. Попробуйте ещё раз:")
        bot.register_next_step_handler(message,
                                       lambda msg: send_task_for_teacher(msg, student_id))
        return
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT id FROM task_list WHERE student_id = ? AND complete IS NULL ", (student_id,))
    tasks = cursor.fetchall()
    if not tasks:
        bot.send_message(message.chat.id, f'У вас нет незавершенных задач.')
        return
    if nomber < 1 or nomber > len(tasks):
        bot.send_message(message.chat.id,
                         f'Введённый номер задачи вне диапазона.')
        return
    tasks_id = tasks[nomber - 1][0]
    bot.send_message(message.chat.id, f"Отправьте файл с решением задачи")
    connection.commit()
    connection.close()
    bot.register_next_step_handler(message, lambda msg: send_document_for_teacher(msg, student_id, tasks_id))


def send_document_for_teacher(message, student_id, tasks_id):
    if message.document:
        file_name = message.document.file_name
        file_info = bot.get_file(message.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        with open(file_name, 'wb') as new_file:
            new_file.write(downloaded_file)
            now = datetime.now()
            current_date = now.strftime("%d.%m")
            current_time = now.strftime("%H:%M")
            connection = sqlite3.connect('my_database.db')
            cursor = connection.cursor()
            cursor.execute("UPDATE task_list SET document = ?, task_time = ?, date = ? WHERE id = ?",
                           (file_name, current_time, current_date, tasks_id))
            connection.commit()
            connection.close()
        bot.send_message(student_id,
                         "Вся информация загружена.\nПосле отправки вам придёт уведомление")
        send_doc_for_teacher(file_name, current_time, current_date, tasks_id, student_id)
    else:
        bot.send_message(message.chat.id, "Пожалуйста, загрузите документ.")
        bot.register_next_step_handler(message,
                                       lambda msg: send_document_for_teacher(msg, student_id, tasks_id))


# ДОСТАЁМ ВСЕ ЗАДАЧИ ИЗ БД ДЛЯ СТУДЕНТА

def get_all_tasks_from_db(message):
    user_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    count_teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    count_student = cursor.fetchone()[0]
    if count_teacher > 0 and count_student > 0:
        delete_zapis(message)
    elif count_teacher > 0 and count_student == 0:
        bot.send_message(message.chat.id, "Функция доступна только студентам.")
    elif count_teacher == 0 and count_student > 0:
        cursor.execute("SELECT task, task_time, date FROM tasks WHERE user_id = ?", (user_id,))
        tasks = cursor.fetchall()
        cursor.execute(
            "SELECT task_id, name_of_discipline, the_task_for_student FROM task_list WHERE student_id = ? AND complete IS NULL ",
            (user_id,))
        tasks_from_teacher = cursor.fetchall()
        output = "".join(f"{i + 1}) {tasks[i][0]} в {tasks[i][1]}, {tasks[i][2]}\n" for i in range(len(tasks)))
        output_task_from_teacher = "".join(
            f"Задача №{tasks_from_teacher[i][0]}\n {tasks_from_teacher[i][1]}\nЗадание:{tasks_from_teacher[i][2]}\n" for i in range(len(tasks_from_teacher)))
        if output or output_task_from_teacher:
            if output:
                bot.send_message(message.chat.id, "Все ваши задачи:")
                bot.send_message(message.chat.id, output)
            if output_task_from_teacher:
                bot.send_message(message.chat.id, output_task_from_teacher)
        else:
            bot.send_message(message.chat.id, "У вас нет задач.")

    else:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "У вас нет учётной записи. Для того, чтобы создать задачу необходимо зарегистрироваться.",
                         reply_markup=markup)
    connection.commit()
    connection.close()

# УДАЛЕНИЕ ЗАДАЧИ ИЗ БД ДЛЯ СТУДЕНТА И ПРЕПОДАВАТЕЛЯ
def delete_task_from_db(message):
    user_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (user_id,))
    count_teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (user_id,))
    count_student = cursor.fetchone()[0]
    if count_teacher > 0 and count_student > 0:
        delete_zapis(message)
    elif count_teacher > 0 and count_student == 0:
        statys = 1
        cursor.execute("SELECT COUNT (*) FROM task_for_student WHERE teacher_id = ?", (user_id,))
        count = cursor.fetchone()[0]
        if count > 0:
            cursor.execute(
                "SELECT id, the_task_for_student, name_of_discipline, send_date, send_time, group_number FROM task_for_student WHERE teacher_id = ?",
                (user_id,))
            tasks = cursor.fetchall()
            proverka_id = "".join(f"{x[0]} " for x in tasks)
            output = "".join(
                f"№{x[0]} - {x[2]}, ЗАДАЧА: {x[1]}\nДЛЯ ГРУППЫ:{x[5]}\nВРЕМЯ ОТПРАВКИ: {x[3]}, {x[4]}\n" for x in tasks)
            bot.send_message(message.chat.id,
                             f"Все ваши задачи:")
            bot.send_message(message.chat.id, output)
            bot.send_message(message.chat.id, "Напишите номер задачи, которую хотите удалить.")
            bot.register_next_step_handler(message, lambda msg: delete_tasks_from_db(msg, proverka_id, statys))
        else:
            bot.send_message(message.chat.id, "У вас нет задач")
    elif count_teacher == 0 and count_student > 0:
        statys = 2
        connection = sqlite3.connect('my_database.db')
        cursor = connection.cursor()
        cursor.execute("SELECT COUNT (*) FROM tasks WHERE user_id = ?", (user_id,))
        count = cursor.fetchone()[0]
        if count > 0:
            cursor.execute("SELECT id, task, task_time, date FROM tasks WHERE user_id = ?", (user_id,))
            tasks = cursor.fetchall()
            proverka_id = "".join(f"{x[0]} " for x in tasks)
            output = "".join(f"{x[0]} - {x[1]} в {x[2]}, {x[3]}\n" for x in tasks)
            bot.send_message(message.chat.id,
                             f"Все ваши задачи:")
            bot.send_message(message.chat.id, output)
            bot.send_message(message.chat.id, "Напишите номер задачи, которую хотите удалить:")
            bot.register_next_step_handler(message, lambda msg: delete_tasks_from_db(msg, proverka_id, statys))
        else:
            bot.send_message(message.chat.id, "У вас нет задач")
    else:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "У вас нет учётной записи. Для того, чтобы создать задачу необходимо зарегистрироваться.",
                         reply_markup=markup)
    connection.commit()
    connection.close()


def delete_tasks_from_db(message, proverka_id, statys):
    id = message.text
    user_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    try:
        if statys == 1:
            if id in proverka_id:
                id = int(id)
                cursor.execute(
                    "SELECT student_id, task_id, name_of_discipline, the_task_for_student FROM task_list WHERE task_id = ?",
                    (id,))
                tasks = cursor.fetchall()
                cursor.execute("DELETE FROM task_for_student WHERE id = ?", (id,))
                cursor.execute("DELETE FROM task_list WHERE task_id = ?", (id,))
                for task in tasks:
                    student_id, task_id, name_of_discipline, the_task_for_student = task
                    bot.send_message(student_id,
                                     f"Задача №{task_id}, {name_of_discipline}\nЗадание:{the_task_for_student}\nУДАЛЕНО ПРЕПОДАВАТЕЛЕМ")
                cursor.execute("SELECT COUNT(*) FROM task_for_student WHERE teacher_id = ?", (user_id,))
                count = cursor.fetchone()[0]
                if count > 0:
                    bot.send_message(message.chat.id, f"Задача {id} удалена")
                else:
                    bot.send_message(message.chat.id, f"Задача {id} удалена")
                    bot.send_message(message.chat.id, "Список задач пуст.")
            else:
                bot.send_message(message.chat.id, "Такого номера нет.")
        elif statys == 2:
            if id in proverka_id:
                id = int(id)
                cursor.execute("DELETE FROM tasks WHERE id = ?", (id,))
                cursor.execute("SELECT COUNT(*) FROM tasks WHERE user_id = ?", (user_id,))
                count = cursor.fetchone()[0]
                if count > 0:
                    bot.send_message(message.chat.id, f"Задача {id} удалена")
                    connection.commit()
                    get_all_tasks_from_db(message)
                else:
                    bot.send_message(message.chat.id, f"Задача {id} удалена")
                    bot.send_message(message.chat.id, "Список задач пуст.")
            else:
                bot.send_message(message.chat.id, "Такого номера нет.")
        connection.commit()
    except Exception as e:
        bot.send_message(message.chat.id, "Произошла ошибка при удалении задачи.")
        print(f"Error: {e}")
    finally:
        connection.close()

def for_admin(message):
    user_id = message.chat.id
    text = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT id_admin FROM parol")
    id_admin = cursor.fetchone()
    id_admin = id_admin[0]
    cursor.execute("INSERT INTO admin (id_user, text_for_admin) VALUES (?, ?)", (user_id, text))
    connection.commit()
    connection.close()
    bot.send_message(user_id, "Обращение зарегистрировано.")
    bot.send_message(id_admin, f"Обращение от пользователя {user_id}\n{text}")

def obrashenya_ot_user(message):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT * FROM admin")
    text = cursor.fetchall()
    connection.commit()
    connection.close()
    if len(text) > 0:
        output = "".join(f"{x[0]} - Пользователь: {x[1]}\n{x[2]}\n" for x in text)
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Отправить ответ на обращение", callback_data="otvet_for_user"))
        bot.send_message(message.chat.id, f"Список обращений:\n{output}", reply_markup=markup)
    else:
        bot.send_message(message.chat.id, "Список обращений пуст.")

def otvet_user(message):
    try:
        nomber = int(message.text)
    except ValueError:
        bot.send_message(message.chat.id,
                         "Неверный номер. Попробуйте ещё раз.")
        bot.register_next_step_handler(message,
                                       lambda msg: obrashenya_ot_user(msg))
        return

    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT id_user FROM admin WHERE id = ?", (nomber,))
    tasks = cursor.fetchone()
    if not tasks:
        bot.send_message(message.chat.id, f'Обращение с таким номером не найдено.')
        return
    user_id = tasks[0]
    bot.send_message(message.chat.id, "Введите текст с ответом:")
    bot.register_next_step_handler(message, lambda msg: delete_text_from_db(msg, user_id, nomber))
    connection.commit()
    connection.close()

def delete_text_from_db(message, user_id, nomber):
    text = message.text
    bot.send_message(user_id, f"Ответ администратора: {text}")
    bot.send_message(message.chat.id, f"Ваш ответ отправлен пользователю.")
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("DELETE FROM admin WHERE id = ?", (nomber,))
    connection.commit()
    connection.close()

#МЕНЯЕМ ПАРОЛЬ (АДМИНИСТРАТОР)
@bot.message_handler(commands=['change_parol'])
def change_parol(message):
    user_id = message.chat.id
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT(*) FROM parol WHERE id_admin = ?", (user_id,))
    count = cursor.fetchone()[0]
    connection.commit()
    connection.close()
    #При смене администратора удалить строчки: "if count >0:" ; "else: bot.send_message(user_id, "Пароль изменяет только администратор.")"
    if count >0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Преподаватели", callback_data="change_parol_teacher"))
        markup.add(types.InlineKeyboardButton("Студенты", callback_data="change_parol_student"))
        bot.send_message(user_id, "Для кого хотите поменять пароль? ", reply_markup=markup)
    else:
        bot.send_message(user_id, "Пароль изменяет только администратор.")

def change_parol_teacher (message, user_id):
    new_parol = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    #ПРИ СМЕНЕ АДМИНА ДОБАВИТЬ СТРОЧКИ:
    #cursor.execute('INSERT INTO parol (id_admin) VALUES (?)',
    #               (user_id,))
    cursor.execute("""UPDATE parol 
                            SET parol_for_teacher = ?
                            WHERE id_admin = ?
                        """, (new_parol, user_id))
    bot.send_message(user_id, "Новый пароль для преподавателей установлен.")
    connection.commit()
    connection.close()

def change_parol_student(message, user_id):
    new_parol = message.text
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("""UPDATE parol 
                            SET parol_for_student = ?
                            WHERE id_admin = ?
                        """, (new_parol, user_id))
    bot.send_message(user_id, "Новый пароль для студентов установлен.")
    connection.commit()
    connection.close()

@bot.message_handler(commands=['manual'])
def settings(message):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    student = cursor.fetchone()[0]
    # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
    if student > 0 and teacher == 0:
        bot.send_message(message.chat.id, f"Инструкция по использованию Телеграмм-бота.\n\n1. Начало работы\nЗапустите бота командой /start и пройдите регистрацию. Чтобы успешно пройти регистрацию "
                                          f"необходим код-пароль, который может сообщить вам администратор. Код необходим для подтверждения вашей роли как студента. После завершения регистрации "
                                          f"вам будут доступны все функции бота. Если вы хотите изменить личные данные - перейдите по пути «Настройки»-«Поменять личные данные»"
                                          f"\n\n2. Создать напоминание\nПерейдите по пути «Мои задачи»-«Создать напоминание», чтобы создать персональное напоминание. Так вы сможете не забыть о важных событиях и делах. Например, вы можете создать напоминание «Выпить стакан воды» на 16:00 (время) 01.12 (дата), напоминание придёт вам в указанное время 1 раз с кнопками «выполнено» и «не выполнено»- они нужны для сбора статистики."
                                          f"\n\n3. Добавить регулярное напоминание\nРегулярное напоминание работает по тому же принцепу, что и обычное, но имеет одно отличие- созданное вами напоминание будет присылаться вам каждый день в одно и то же время, которое вы указали. Например, если вы создаёте регулярное напоминание «Выпить стакан воды» на 16:00 (время) 01.12 (дата), то начиная с 01.12 бот будет отправлять вам напоминание каждый деньв 16:00 "
                                          f"\n\n4. Все задачи\nПерейдите по пути «Мои задачи»-«Все мои задачи» и бот пришлёт вам все персональные регулярные и нерегулярные напоминания, которые вы создавали. Также бот пришлёт вам список задач от преподавателей, которые вами не выполнены, если такие имеются."
                                          f"\n\n5. Отправить решение преподавателю\nЕсли у вас есть невыполненные задачи от преподавателя, вы сможете отправить решение если перейдете по пути «Мои задачи»-«Отправить решение преподавателю». Там вам нужно будет выбрать задание, по которому хотите прикрепите решение (в формате документа). После того как вы загрузите документ, он будет отправлен преподавателю. Далее преподаватель сможет оценить вашу работу и прислать вам свой комментарий по выполненной работе с оценкой. "
                                          f"\n\n6. Удалить напоминание \nПерейдите по пути «Мои задачи»-«Удалить напоминание» для того, чтобы удалить персональное напоминание."
                                          f"\n\n7. Статистика\nВы сможете увидеть статистику, если введёте команду /statistics. Вам придёт круговая диаграмма, на которой будут отражены выплненные и невыполненные вами задачи. "
                                          f"\n\n8. Удалить аккаунт\nПерейдите по пути «Настройки»-«Удалить аккаунт» для того, чтобы удалить все данные о себе. Чтобы начать работать с ботом вам прийдется снова пройти регистрацию. Ваши данные полностью будут удалены."
                                          f"\n\n9. Обращение к администратору\nЕсли бот вам выдаёт ошибку или некорректно выводит данные, то вы можете обратиться за помощью к администратору, если перейдёте поп пути «Настройки»-«Обращение к администратору»"
                                           )

    elif student == 0 and teacher > 0:
        bot.send_message(message.chat.id, f"Инструкция по использованию Телеграмм-бота.\n\n1. Начало работы\nЗапустите бота командой /start и пройдите регистрацию. Чтобы успешно пройти регистрацию "
                                          f"необходим код-пароль, который может сообщить вам администратор. Код необходим для подтверждения вашей роли как преподавателя. После завершения регистрации "
                                          f"вам будут доступны все функции бота. Если вы хотите изменить личные данные - перейдите по пути «Настройки»-«Поменять личные данные»"
                                          f"\n\n2. Создание дисциплины и группы\nВведите команду /settings для перехода в настройки или перейдите в «меню» и выберите «настройки». Чтобы вы смогли "
                                          f"отправлять задачи студентом необходимо создать группу и дисциплину, в которые студенты смогут добавляться при регистрации. Для создания новой дисциплины нажмите "
                                          f"«создать дисциплину». Для создания группы нажмите «Добавить группу». Вы также можете изменять названия, если перейдете по пути «Настройки»-«все дисциплины»-«изменить "
                                          f"наименование», для дисциплины и «Настройки»-«все группы»-«изменить группу», для группы.\n\n"
                                          f"3. Отправка задания\nЧтобы отправить задачу студентам перейдите по пути «Мои задачи»-«Создать задачу», после чего необходимо выбрать группу, которой будет отправлена "
                                          f"задача. Задача будет отправлена каждому участнику группы, который зарегистрирован в ней. Прикрепите файл с вашим заданием (Документ любого формата) и выберите время, в которое "
                                          f"необходимо осуществить отправку. После этого ваше задание будет отправлено студентам, и вам придёт уведомление об успешной отправке.\n\n"
                                          f"4. Отправка оценки и комментария по выполненной работе\nЧтобы оценить работу и отправить свой комментарий студенту необходимо перейти по пути «Мои задачи»-«Оценить работу». Далее выберите номер решения от студента, по которому хотите отправить оценку, и введите оценку и ваш комментарий по работе. После этого вся информация будет отправлена студенту, и оценка будет отражена в отчётном документе. Вы сможете это увидеть в отделе «Статистика»."
                                          f"\n\n5. Отправить уведомление студентам\nЕсли у вас есть важная информация, которую необходимо донести до всех студентов группы, то вы сможете её отправить если перейдете по пути «Мои задачи»-«Отправить уведомление студентам». Там вам нужно будет выбрать группу, которой будет отправлено ваше сообщение. Далее вы вводите текст, к примеру, «Дорогие студенты, зачет будет проходить завтра в 10:00 в аудитории 256.», он будет отправлен всем студентам. "
                                          f"\n\n6. Информация о студентах\nВы можете узнать номер телефона и почту студента, чтобы связаться с ним напрямую. Необходимо перейти по пути «Настройки»-«Информация о студентах». Вам будут доступны контактные данные, которые указывал студент при регистрации."
                                          f"\n\n7. Статистика\nВы сможете увидеть статистику, если введёте команду /statistics. На выбор будет два варианта: где 1 - краткая сводка о количестве выполненных и невыполненных задач, и 2 - Отчёт, в формате документа word, за указанный вами временной промежуток, где будет информация по каждому студенту группы, количество выполненных работ, все оценки и средний балл. "
                                          f"\n\n8. Удаление задачи\nПерейдите по пути «Мои задачи»-«Удалить задачу» для того, чтобы удалить задачу. После удаления, об этом придёт уведомление всем студентам, которым эта задача была отправлена."
                                          f"\n\n9. Удалить аккаунт\nПерейдите по пути «Настройки»-«Удалить аккаунт» для того, чтобы удалить все данные о себе. Чтобы начать работать с ботом вам прийдется снова пройти регистрацию. Ваши данные полностью будут удалены."
                                          f"\n\n10. Обращение к администратору\nЕсли бот вам выдаёт ошибку или некорректно выводит данные, то вы можете обратиться за помощью к администратору, если перейдёте поп пути «Настройки»-«Обращение к администратору»")

    elif student > 0 and teacher > 0:
        bot.send_message(message.chat.id,
                         f"Вы зарегистрированы и как преподаватель и как студент. Невозможно продолжить работу.")
    elif student == 0 and teacher == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "Вы не зарегистрированы. Для того, чтобы продолжить, необходимо пройти регистрацию. ",
                         reply_markup=markup)
    connection.commit()
    connection.close()

@bot.message_handler(commands=['settings'])
def settings(message):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    student = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM parol WHERE id_admin = ?", (message.chat.id,))
    admin = cursor.fetchone()[0]
    # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
    if student > 0 and teacher == 0:
        if admin > 0:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Изменить пароль (администратор)", callback_data="change_parol"))
            markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
            markup.add(types.InlineKeyboardButton("Поменять личные данные", callback_data="changing_settings_student"))
            markup.add(types.InlineKeyboardButton("Удалить аккаунт", callback_data="delete_account"))
            markup.add(types.InlineKeyboardButton("Обращение к администратору", callback_data="text_for_admin"))
            bot.send_message(message.chat.id, "Выберите кнопку:", reply_markup = markup)
        else:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
            markup.add(types.InlineKeyboardButton("Поменять личные данные", callback_data="changing_settings_student"))
            markup.add(types.InlineKeyboardButton("Удалить аккаунт", callback_data="delete_account"))
            markup.add(types.InlineKeyboardButton("Обращение к администратору", callback_data="text_for_admin"))
            bot.send_message(message.chat.id, "Выберите кнопку:", reply_markup=markup)
    elif student == 0 and teacher > 0:
        if admin > 0:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Изменить пароль (администратор)", callback_data="change_parol"))
            markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
            markup.add(types.InlineKeyboardButton("Поменять личные данные", callback_data="changing_settings_teacher"))
            markup.add(types.InlineKeyboardButton("Удалить аккаунт", callback_data="delete_account"))
            markup.add(types.InlineKeyboardButton("Создать дисциплину", callback_data="add"))
            markup.add(types.InlineKeyboardButton("Все дисциплины", callback_data="all_discipline"))
            markup.add(types.InlineKeyboardButton("Создать группу", callback_data="add_group"))
            markup.add(types.InlineKeyboardButton("Все группы", callback_data="all_grupp"))
            markup.add(types.InlineKeyboardButton("Информация о студентах", callback_data="info_about_student"))
            markup.add(types.InlineKeyboardButton("Обращение к администратору", callback_data="text_for_admin"))
            bot.send_message(message.chat.id, "Выберите кнопку:", reply_markup=markup)
        else:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
            markup.add(types.InlineKeyboardButton("Поменять личные данные", callback_data="changing_settings_teacher"))
            markup.add(types.InlineKeyboardButton("Удалить аккаунт", callback_data="delete_account"))
            markup.add(types.InlineKeyboardButton("Создать дисциплину", callback_data="add"))
            markup.add(types.InlineKeyboardButton("Все дисциплины", callback_data="all_discipline"))
            markup.add(types.InlineKeyboardButton("Создать группу", callback_data="add_group"))
            markup.add(types.InlineKeyboardButton("Все группы", callback_data="all_grupp"))
            markup.add(types.InlineKeyboardButton("Информация о студентах", callback_data="info_about_student"))
            markup.add(types.InlineKeyboardButton("Обращение к администратору", callback_data="text_for_admin"))
            bot.send_message(message.chat.id, "Выберите кнопку:", reply_markup=markup)
    elif student > 0 and teacher > 0:
        bot.send_message(message.chat.id,
                         f"Вы зарегистрированы и как преподаватель и как студент. Невозможно продолжить работу.")
    elif student == 0 and teacher == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "Вы не зарегистрированы. Для того, чтобы продолжить, необходимо пройти регистрацию. ",
                         reply_markup=markup)
    connection.commit()
    connection.close()

@bot.message_handler(commands=['my_tasks'])
def settings(message):
    connection = sqlite3.connect('my_database.db')
    cursor = connection.cursor()
    cursor.execute("SELECT COUNT (*) FROM teachers WHERE teacher_id = ?", (message.chat.id,))
    teacher = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM student WHERE student_id = ?", (message.chat.id,))
    student = cursor.fetchone()[0]
    cursor.execute("SELECT COUNT (*) FROM parol WHERE id_admin = ?", (message.chat.id,))
    admin = cursor.fetchone()[0]
    # ПРОВЕРКА ЗАРЕГ-Н ПОЛЬЗОВАТЕЛЬ ИЛИ НЕТ
    if student > 0 and teacher == 0:
        if admin > 0:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Обращения от пользователей", callback_data="user_text"))
            markup.add(types.InlineKeyboardButton("Создать напоминание", callback_data="add_task"))
            markup.add(types.InlineKeyboardButton("Добавить регулярное напоминание", callback_data="add_regular_task"))
            markup.add(types.InlineKeyboardButton("Все мои задачи", callback_data="all_tasks"))
            markup.add(
                types.InlineKeyboardButton("Отправить решение преподавателю", callback_data="task_from_the_teacher"))
            markup.add(types.InlineKeyboardButton("Удалить напоминание", callback_data="delete_tasks"))
            bot.send_message(message.chat.id, "Выберите кнопку:", reply_markup=markup)
        else:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Создать напоминание", callback_data="add_task"))
            markup.add(types.InlineKeyboardButton("Добавить регулярное напоминание", callback_data="add_regular_task"))
            markup.add(types.InlineKeyboardButton("Все мои задачи", callback_data="all_tasks"))
            markup.add(types.InlineKeyboardButton("Отправить решение преподавателю", callback_data="task_from_the_teacher"))
            markup.add(types.InlineKeyboardButton("Удалить напоминание", callback_data="delete_tasks"))
            bot.send_message(message.chat.id, "Выберите кнопку:", reply_markup=markup)
    elif student == 0 and teacher > 0:
        if admin >0:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Обращения от пользователей", callback_data="user_text"))
            markup.add(types.InlineKeyboardButton("Создать задачу", callback_data="add_task"))
            markup.add(types.InlineKeyboardButton("Удалить задачу (студентам придёт уведомление об удалении)",
                                                  callback_data="delete_tasks"))
            markup.add(types.InlineKeyboardButton("Оценить работу студента", callback_data="send_markk"))
            markup.add(
                types.InlineKeyboardButton("Отправить уведомление студентам", callback_data="send_message_for_student"))
            bot.send_message(message.chat.id, "Выберите кнопку:", reply_markup=markup)
        else:
            markup = types.InlineKeyboardMarkup()
            markup.add(types.InlineKeyboardButton("Создать задачу", callback_data="add_task"))
            markup.add(types.InlineKeyboardButton("Удалить задачу (студентам придёт уведомление об удалении)", callback_data="delete_tasks"))
            markup.add(types.InlineKeyboardButton("Оценить работу студента", callback_data="send_markk"))
            markup.add(types.InlineKeyboardButton("Отправить уведомление студентам", callback_data="send_message_for_student"))
            bot.send_message(message.chat.id, "Выберите кнопку:", reply_markup=markup)
    elif student > 0 and teacher > 0:
        bot.send_message(message.chat.id,
                         f"Вы зарегистрированы и как преподаватель и как студент. Невозможно продолжить работу.")
    elif student == 0 and teacher == 0:
        markup = types.InlineKeyboardMarkup()
        markup.add(types.InlineKeyboardButton("Регистрация", callback_data="registration"))
        bot.send_message(message.chat.id,
                         "Вы не зарегистрированы. Для того, чтобы продолжить, необходимо пройти регистрацию. ",
                         reply_markup=markup)
    connection.commit()
    connection.close()
# НАПОМИНАНИЕ ПОЛЬЗОВАТЕЛЮ (Преподаватель)
def send_message_ga(user_id, message):
    bot.send_message(chat_id=user_id, text=f"{message}")

# НАПОМИНАНИЕ ПОЛЬЗОВАТЕЛЮ (Студент)
def send_message_ga_student(user_id, message):
    markup = types.InlineKeyboardMarkup()
    markup.add(types.InlineKeyboardButton("Выполнено", callback_data="done"))
    markup.add(types.InlineKeyboardButton("Не выполнено", callback_data="dont_done"))
    bot.send_message(chat_id=user_id, text=f"НАПОМИНАНИЕ: {message}", reply_markup=markup)

# ОТПРАВЛЕНИЕ ПЕРСОНАЛЬНОЙ ЗАДАЧИ СТУДЕНТА
def check_tasks():
    conn = sqlite3.connect('my_database.db')
    cursor = conn.cursor()
    now = datetime.now()
    current_time = now.strftime("%H:%M")
    current_date = now.strftime("%d.%m")
    # Выборка задач с текущим временем и датой
    cursor.execute("""
        SELECT id, user_id, task, regular_task 
        FROM tasks 
        WHERE task_time = ? AND date = ?
    """, (current_time, current_date))
    tasks = cursor.fetchall()
    for task in tasks:
        task_id, user_id, message, regular_task = task
        send_message_ga_student(user_id, message)
        if regular_task == 1:
            next_day = now + timedelta(days=1)
            current_datee = next_day.strftime("%d.%m")
            cursor.execute("""
                            UPDATE tasks 
                            SET date = ?
                            WHERE id = ?
                        """, (current_datee, task_id))
            cursor.execute("""UPDATE tasks 
                            SET count_regular_task = count_regular_task + 1
                            WHERE id = ?
                        """, (task_id,))
            cursor.execute("""                
                            UPDATE statystic_for_student 
                            SET all_tasks = all_tasks + 1, 
                                dont_complete = dont_complete + 1
                            WHERE student_id = ?""", (user_id,))
        # Если задача не регулярная, обновляем complete на True
        if not regular_task:
            cursor.execute("""
                UPDATE tasks 
                SET complete = 1 
                WHERE id = ?
            """, (task_id,))
            cursor.execute("""                
                UPDATE statystic_for_student 
                SET all_tasks = all_tasks + 1, 
                    dont_complete = dont_complete + 1
                WHERE student_id = ?""", (user_id,))

    conn.commit()
    conn.close()


# ОТПРАВЛЕНИЕ ЗАДАЧИ ПРЕПОДАВАТЕЛЯ
def send_doc():
    conn = sqlite3.connect('my_database.db')
    cursor = conn.cursor()
    now = datetime.now()
    current_time = now.strftime("%H:%M")
    current_date = now.strftime("%d.%m")
    cursor.execute("""
            SELECT id, document, group_number, name_of_discipline, the_task_for_student, teacher_id, statys
            FROM task_for_student 
            WHERE send_time = ? AND send_date = ? """, (current_time, current_date))
    tasks = cursor.fetchall()
    for task in tasks:
        id, document, group_number, name_of_discipline, the_task_for_student, teacher_id, statys = task
        if not statys:
            cursor.execute("""
                UPDATE task_for_student 
                SET statys = 1 
                WHERE id = ?
            """, (id,))
            cursor.execute("""
                    SELECT student_id 
                    FROM student 
                    WHERE group_number = ? """, (group_number,))
            for_student = cursor.fetchall()
            for student in for_student:
                student_id = student[0]
                send_message_ga(student_id, f"{name_of_discipline}\nЗАДАНИЕ:\n{the_task_for_student}\n")
                bot.send_document(student_id, open(f"{document}", "rb"))
                cursor.execute("""
                                        SELECT name 
                                        FROM student 
                                        WHERE student_id = ? """, (student_id,))
                name_student = cursor.fetchall()
                for names in name_student:
                    student_name = names[0]
                    cursor.execute(
                        'INSERT INTO task_list (name_student, task_id, student_id, teacher_id, name_of_discipline, the_task_for_student, group_number, send_teacher_for_student_date) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                        (student_name, id, student_id, teacher_id, name_of_discipline, the_task_for_student, group_number, current_date))
            send_message_ga(teacher_id,
                            f"{name_of_discipline}\nЗадача: {the_task_for_student}\nотправлена студентам группы {group_number}")
            cursor.execute("UPDATE task_for_student SET document = NULL WHERE id= ?", (id,))
        os.remove(f"{document}")
    conn.commit()
    conn.close()


# ОТПРАВКА РЕШЕНЁННОГО ЗАДАНИЯ ПРЕПОДАВАТЕЛЮ
def send_doc_for_teacher(file_name, current_time, current_date, tasks_id, student_id):
    conn = sqlite3.connect('my_database.db')
    cursor = conn.cursor()
    cursor.execute("""
        UPDATE task_list 
        SET complete = 1 
        WHERE document = ? AND task_time = ? AND date = ? AND id = ?
    """, (file_name, current_time, current_date, tasks_id))
    cursor.execute("""
            SELECT name 
            FROM student 
            WHERE student_id = ? """, (student_id,))
    name_student = cursor.fetchall()
    for name in name_student:
        student_name = name[0]
        cursor.execute("""
                    SELECT teacher_id, name_of_discipline, the_task_for_student, group_number
                    FROM task_list 
                    WHERE student_id = ? AND document = ? AND task_time = ? AND date = ? AND id = ? """, (student_id, file_name, current_time, current_date, tasks_id))
        name_te = cursor.fetchall()
        for n in name_te:
            teacher_id, name_of_discipline, the_task_for_student, group_number = n
            send_message_ga(teacher_id,
                            f"Решение задачи №{tasks_id}\n{name_of_discipline}\nЗАДАНИЕ:\n{the_task_for_student}\nОт студента: {student_name} группа {group_number} ")
            bot.send_document(teacher_id, open(f"{file_name}", "rb"))
            cursor.execute("UPDATE task_list SET name_student = ? WHERE student_id = ? AND document = ? AND task_time = ? AND date = ? AND id = ?", (student_name, student_id, file_name, current_time, current_date, tasks_id))
            send_message_ga(student_id,
                            f"{name_of_discipline}\nРешение задачи: {the_task_for_student}\n отправлено.")

    conn.commit()
    conn.close()
    os.remove(f"{file_name}")

def send_coment_teacher(current_date,current_time):
    conn = sqlite3.connect('my_database.db')
    cursor = conn.cursor()
    cursor.execute("""
            SELECT task_id, student_id, teacher_id, name_of_discipline, the_task_for_student, send_time, date, group_number, comment, mark
            FROM teacher_comment 
            WHERE send_time = ? AND date = ? """, (current_time, current_date))
    comments = cursor.fetchall()
    for comm in comments:
        task_id, student_id, teacher_id, name_of_discipline, the_task_for_student, send_time, date, group_number, comment, mark = comm
        cursor.execute("""
                SELECT name 
                FROM teachers 
                WHERE teacher_id = ? """, (teacher_id,))
        name_teacher = cursor.fetchall()
        for name in name_teacher:
            teacher_name = name[0]
            send_message_ga(student_id,
                            f"КОММЕНТАРИЙ ПРЕПОДАВАТЕЛЯ:\n{teacher_name}\nОценка {mark} по №{task_id}\n{name_of_discipline}\nЗАДАНИЕ:\n{the_task_for_student}\nКомментарий: {comment} ")
        send_message_ga(teacher_id,
                        f"Комментарий отправлен.")
        cursor.execute("UPDATE task_list SET send_mark_date = ? WHERE task_id = ?", (current_date, task_id))
    conn.commit()
    conn.close()

def reset_tasks():
    try:
        conn = sqlite3.connect('my_database.db')
        cursor = conn.cursor()
        cursor.execute("DELETE FROM statystic_for_student")
        conn.commit()
        conn.close()
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")

scheduler = BackgroundScheduler()
# Запланируем выполнение функции check_tasks каждую минуту
scheduler.add_job(check_tasks, 'interval', minutes=1)
scheduler.add_job(send_doc, 'interval', minutes=1)
scheduler.add_job(reset_tasks, 'cron', day=1, hour=0, minute=0)
scheduler.start()

bot.polling(none_stop=True)



